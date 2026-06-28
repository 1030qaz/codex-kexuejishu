#!/usr/bin/env python3
"""Build weekly multi-user original-post and analyzed documents."""

from __future__ import annotations

import argparse
import json
from collections import defaultdict
from dataclasses import asdict
from datetime import datetime, timedelta
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from build_june_analyzed_doc import StockResolver, build_doc, fetch_index_klines, load_posts, write_markdown as write_analysis_markdown
from build_multiuser_monthly_docs import THREAD_TITLE, UserPost, add_post, sorted_posts, style_document


def parse_dt(value: str) -> datetime:
    return datetime.fromisoformat(value)


def week_start(day: datetime) -> datetime:
    return (day - timedelta(days=day.weekday())).replace(hour=0, minute=0, second=0, microsecond=0)


def week_end(start: datetime) -> datetime:
    return (start + timedelta(days=6)).replace(hour=23, minute=59, second=59, microsecond=0)


def week_key(start: datetime) -> str:
    iso = start.isocalendar()
    return f"{iso.year}-W{iso.week:02d}"


def load_user_posts(path: Path, start_date: str, end_date: str | None) -> list[UserPost]:
    data = json.loads(path.read_text(encoding="utf-8-sig"))
    posts: list[UserPost] = []
    for item in data:
        day = item.get("posted_at", "")[:10]
        if day < start_date:
            continue
        if end_date and day > end_date:
            continue
        posts.append(UserPost(**item))
    return sorted_posts(posts)


def group_weeks(posts: list[UserPost]) -> dict[str, tuple[datetime, datetime, list[UserPost]]]:
    grouped: dict[str, list[UserPost]] = defaultdict(list)
    starts: dict[str, datetime] = {}
    for post in posts:
        start = week_start(parse_dt(post.posted_at))
        key = week_key(start)
        grouped[key].append(post)
        starts[key] = start
    return {key: (starts[key], week_end(starts[key]), sorted_posts(items)) for key, items in sorted(grouped.items())}


def build_week_doc(label: str, start: datetime, end: datetime, posts: list[UserPost], output_dir: Path) -> Path:
    document = Document()
    style_document(document)
    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(f"{THREAD_TITLE}：{label} 发言整理")
    summary = document.add_paragraph()
    summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    wolf_count = sum(1 for post in posts if post.uid == "150058")
    summary.add_run(f"{start:%Y-%m-%d} 至 {end:%Y-%m-%d}；共 {len(posts)} 条；狼大 {wolf_count} 条")
    for post in posts:
        add_post(document, post)
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / f"{THREAD_TITLE}_发言整理_{label}.docx"
    document.save(output_path)
    return output_path


def write_week_markdown(label: str, start: datetime, end: datetime, posts: list[UserPost], output_dir: Path) -> Path:
    lines = [f"# {THREAD_TITLE}：{label} 发言整理", "", f"> {start:%Y-%m-%d} 至 {end:%Y-%m-%d}；共 {len(posts)} 条", ""]
    for post in posts:
        lines.append(f"## {post.time} | {post.floor}楼 | {post.username} | UID {post.uid}")
        lines.append("")
        lines.append(post.body)
        lines.append("")
    output_path = output_dir / f"{THREAD_TITLE}_发言整理_{label}.md"
    output_path.write_text("\n".join(lines), encoding="utf-8")
    return output_path


def build_week_analysis(
    combined_json: Path,
    label: str,
    start: datetime,
    end: datetime,
    output_dir: Path,
    resolver: StockResolver,
) -> tuple[Path, Path, int]:
    start_day = start.strftime("%Y-%m-%d")
    end_day = end.strftime("%Y-%m-%d")
    posts = load_posts(combined_json, "all", start_day, end_day)
    if not posts:
        raise ValueError(f"No posts found for {label}")

    cache_path = Path(f"market_cache_{label}.json")
    market_days = fetch_index_klines(start.strftime("%Y%m%d"), end.strftime("%Y%m%d"), cache_path)
    analysis_doc = output_dir / f"{THREAD_TITLE}_发言逐条分析_{label}.docx"
    analysis_md = output_dir / f"{THREAD_TITLE}_发言逐条分析_{label}.md"
    report_label = f"{label}（{start_day} 至 {end_day}）"
    build_doc(posts, market_days, resolver, analysis_doc, report_label)
    write_analysis_markdown(posts, market_days, resolver, analysis_md, report_label)
    return analysis_doc, analysis_md, len(posts)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build weekly docs and analysis for selected users since 2026.")
    parser.add_argument("--combined-json", default="selected_users_posts.json")
    parser.add_argument("--output-dir", default="weekly_docs")
    parser.add_argument("--start-date", default="2026-01-01")
    parser.add_argument("--end-date", help="Optional final date, defaults to latest post date in JSON.")
    parser.add_argument("--stock-cache", default="stock_name_cache.json")
    parser.add_argument("--only-week", help="Optional ISO week key, e.g. 2026-W26.")
    parser.add_argument("--skip-original", action="store_true", help="Only build analyzed weekly documents.")
    parser.add_argument("--skip-analysis", action="store_true", help="Only build original weekly documents.")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    combined_json = Path(args.combined_json)
    output_dir = Path(args.output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    raw_posts = load_user_posts(combined_json, args.start_date, args.end_date)
    if not raw_posts:
        print("No posts found.")
        return 1

    weeks = group_weeks(raw_posts)
    if args.only_week:
        weeks = {key: value for key, value in weeks.items() if key == args.only_week}
    if not weeks:
        print("No matching weeks found.")
        return 1

    resolver = StockResolver(Path(args.stock_cache))
    results = []
    for label, (start, end, posts) in weeks.items():
        week_result: dict[str, object] = {
            "week": label,
            "start": start.strftime("%Y-%m-%d"),
            "end": end.strftime("%Y-%m-%d"),
            "posts": len(posts),
        }
        if not args.skip_original:
            week_doc = build_week_doc(label, start, end, posts, output_dir)
            week_md = write_week_markdown(label, start, end, posts, output_dir)
            week_result["week_doc"] = str(week_doc)
            week_result["week_markdown"] = str(week_md)
        if not args.skip_analysis:
            analysis_doc, analysis_md, count = build_week_analysis(combined_json, label, start, end, output_dir, resolver)
            week_result["analysis_doc"] = str(analysis_doc)
            week_result["analysis_markdown"] = str(analysis_md)
            week_result["analysis_posts"] = count
        results.append(week_result)
        print(json.dumps(week_result, ensure_ascii=False), flush=True)

    resolver.save()
    summary_path = output_dir / "weekly_build_summary_2026.json"
    summary_path.write_text(json.dumps(results, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    print(json.dumps({"weeks": len(results), "summary": str(summary_path)}, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
