from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE = Path(__file__).resolve().parent
MD_PATH = BASE / "A股新能源方向学习与分析手册.md"
DOCX_PATH = BASE / "A股新能源方向学习与分析手册.docx"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(run, name="Microsoft YaHei", size=None, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_font(paragraph, name="Microsoft YaHei", size=10.5):
    for run in paragraph.runs:
        set_run_font(run, name=name, size=size)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("bottom", bottom), ("start", start), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for w in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.tcW
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[idx]))
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for side in ("top_margin", "right_margin", "bottom_margin", "left_margin"):
        setattr(section, side, Inches(1))
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        st = styles[style_name]
        st.font.name = "Microsoft YaHei"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.line_spacing = 1.25


def add_footer(doc):
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("A股新能源方向学习与分析手册 | 学习研究用途，不构成投资建议")
        set_run_font(run, size=8.5, color="666666")


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(80)
    title = p.add_run("A股新能源方向学习与分析手册")
    set_run_font(title, size=24, bold=True, color="0B2545")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = p.add_run("产业链研究 · 周期判断 · 公司池跟踪 · 交易复盘")
    set_run_font(subtitle, size=13, color="1F4D78")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta = p.add_run("版本：2026-06-25 | 数据口径：实时字段均需更新验证")
    set_run_font(meta, size=10, color="555555")

    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    widths = [2200, CONTENT_WIDTH_DXA - 2200]
    set_table_geometry(table, widths)
    items = [
        ("读者定位", "股票小白到主题研究入门者。"),
        ("核心目标", "把新能源方向拆成可验证、可更新、可复盘的研究框架。"),
        ("数据原则", "东方财富接口优先；不可用时用腾讯财经、交易所、巨潮和行业数据兜底。"),
        ("重要声明", "仅用于学习、研究和跟踪，不提供买卖建议。"),
    ]
    for row, (k, v) in zip(table.rows, items):
        row.cells[0].text = k
        row.cells[1].text = v
        set_cell_shading(row.cells[0], "E8EEF5")
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, size=9.5)
    doc.add_page_break()


def collect_toc(markdown: str):
    items = []
    for line in markdown.splitlines():
        m = re.match(r"^(#{2,3})\s+(.+)", line)
        if m:
            title = m.group(2).strip()
            if title.startswith("附录") or title.startswith("第"):
                items.append((len(m.group(1)), title))
    return items


def add_static_toc(doc, markdown):
    doc.add_heading("目录", level=1)
    for level, title in collect_toc(markdown)[:80]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25 if level == 3 else 0)
        run = p.add_run(title)
        set_run_font(run, size=9.5 if level == 3 else 10.5, bold=(level == 2))
    doc.add_page_break()


def add_callout(doc, text, label="提示"):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(f"{label}：")
    set_run_font(r, size=10, bold=True, color="1F3A5F")
    r = p.add_run(text)
    set_run_font(r, size=10)


def is_table_line(line):
    return line.strip().startswith("|") and line.strip().endswith("|")


def parse_table(lines, start):
    rows = []
    i = start
    while i < len(lines) and is_table_line(lines[i]):
        raw = lines[i].strip().strip("|")
        cells = [c.strip() for c in raw.split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in cells):
            rows.append(cells)
        i += 1
    return rows, i


def widths_for_table(col_count):
    if col_count == 1:
        return [CONTENT_WIDTH_DXA]
    if col_count == 2:
        return [2200, CONTENT_WIDTH_DXA - 2200]
    if col_count == 3:
        return [1800, 3600, CONTENT_WIDTH_DXA - 5400]
    if col_count == 4:
        return [1600, 2600, 2600, CONTENT_WIDTH_DXA - 6800]
    if col_count == 5:
        return [1500, 1900, 2100, 1900, CONTENT_WIDTH_DXA - 7400]
    if col_count == 6:
        return [1200, 1700, 1700, 1700, 1500, CONTENT_WIDTH_DXA - 7800]
    return [int(CONTENT_WIDTH_DXA / col_count)] * col_count


def add_markdown_table(doc, rows):
    if not rows:
        return
    max_cols = max(len(r) for r in rows)
    rows = [r + [""] * (max_cols - len(r)) for r in rows]

    if max_cols > 7:
        add_callout(doc, "下列表格为宽字段模板。为保证Word可读性，DOCX版转换为字段清单；Markdown源文件保留完整宽表。", "宽表说明")
        header = rows[0]
        for data in rows[1:]:
            p = doc.add_paragraph(style=None)
            p.paragraph_format.left_indent = Inches(0.15)
            text = "；".join(f"{h}: {v}" for h, v in zip(header, data) if h or v)
            r = p.add_run(text)
            set_run_font(r, size=9)
        return

    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = "Table Grid"
    set_table_geometry(table, widths_for_table(max_cols))
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = value
            if r_idx == 0:
                set_cell_shading(cell, "E8EEF5")
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    set_run_font(run, size=8.5 if max_cols >= 5 else 9.3, bold=(r_idx == 0))
    doc.add_paragraph()


def add_code_block(doc, code):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F7F7F7")
    cell.text = ""
    for idx, line in enumerate(code.rstrip("\n").splitlines()):
        p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        r = p.add_run(line)
        set_run_font(r, name="Consolas", size=8.5, color="333333")


def add_markdown_paragraph(doc, text):
    stripped = text.strip()
    if not stripped:
        return
    if stripped.startswith(">"):
        add_callout(doc, stripped.lstrip("> ").strip(), "说明")
        return
    if stripped.startswith("- "):
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(stripped[2:].strip())
        set_run_font(r, size=10)
        return
    if re.match(r"^\d+\.\s+", stripped):
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(re.sub(r"^\d+\.\s+", "", stripped))
        set_run_font(r, size=10)
        return
    p = doc.add_paragraph()
    r = p.add_run(stripped.replace("`", ""))
    set_run_font(r, size=10.5)


def build():
    markdown = MD_PATH.read_text(encoding="utf-8")
    doc = Document()
    style_document(doc)
    add_cover(doc)
    add_static_toc(doc, markdown)

    lines = markdown.splitlines()
    i = 0
    in_code = False
    code_lines = []
    while i < len(lines):
        line = lines[i]
        if line.strip().startswith("```"):
            if in_code:
                add_code_block(doc, "\n".join(code_lines))
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue
        if is_table_line(line):
            rows, i = parse_table(lines, i)
            add_markdown_table(doc, rows)
            continue
        m = re.match(r"^(#{1,6})\s+(.+)", line)
        if m:
            level = min(len(m.group(1)), 3)
            title = m.group(2).strip()
            if level == 1:
                doc.add_heading(title, level=1)
            else:
                doc.add_heading(title, level=level - 1)
        else:
            add_markdown_paragraph(doc, line)
        i += 1

    add_footer(doc)
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build()
