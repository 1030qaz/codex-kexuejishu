#!/usr/bin/env python3
"""Update tracked NGA users and rebuild one weekly report set."""

from __future__ import annotations

import argparse
import json
import os
import subprocess
import sys
from dataclasses import asdict
from datetime import datetime, timedelta
from pathlib import Path

from build_multiuser_monthly_docs import THREAD_TITLE, UserPost, add_post, sorted_posts, style_document
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


def week_bounds(value: str | None = None) -> tuple[datetime, datetime]:
    if value:
        day = datetime.strptime(value, "%Y-%m-%d")
    else:
        day = datetime.now()
    start = day - timedelta(days=day.weekday())
    end = start + timedelta(days=6)
    return start.replace(hour=0, minute=0, second=0, microsecond=0), end.replace(hour=23, minute=59, second=59, microsecond=0)


def week_key(start: datetime) -> str:
    iso = start.isocalendar()
    return f"{iso.year}-W{iso.week:02d}"


def load_posts(path: Path) -> list[UserPost]:
    data = json.loads(path.read_text(encoding="utf-8-sig"))
    return sorted_posts([UserPost(**item) for item in data])


def filter_week(posts: list[UserPost], start: datetime, end: datetime) -> list[UserPost]:
    result: list[UserPost] = []
    for post in posts:
        posted_at = datetime.fromisoformat(post.posted_at)
        if start <= posted_at <= end:
            result.append(post)
    return sorted_posts(result)


def run_step(command: list[str]) -> None:
    print("Running: " + " ".join(command), flush=True)
    subprocess.run(command, check=True)


def build_week_doc(label: str, start: datetime, end: datetime, posts: list[UserPost], output_dir: Path) -> Path:
    document = Document()
    style_document(document)
    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(f"{THREAD_TITLE}：{label} 发言整理")
    summary = document.add_paragraph()
    summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    summary.add_run(f"{start:%Y-%m-%d} 至 {end:%Y-%m-%d}；共 {len(posts)} 条")
    for post in posts:
        add_post(document, post)
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / f"{THREAD_TITLE}_发言整理_{label}.docx"
    document.save(output_path)
    return output_path


def write_week_markdown(label: str, start: datetime, end: datetime, posts: list[UserPost], output_dir: Path) -> Path:
    lines = [f"# {THREAD_TITLE}：{label} 发言整理", "", f"> {start:%Y-%m-%d} 至 {end:%Y-%m-%d}；共 {len(posts)} 条", ""]
    for post in posts:
        lines.append(f"## {post.time} | {post.floor}楼 | {post.username} | UID {post.uid}")
        lines.append("")
        lines.append(post.body)
        lines.append("")
    output_path = output_dir / f"{THREAD_TITLE}_发言整理_{label}.md"
    output_path.write_text("\n".join(lines), encoding="utf-8")
    return output_path


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Update users and rebuild one weekly report set.")
    parser.add_argument("--week-of", help="Any date inside the target week, e.g. 2026-06-22. Defaults to today.")
    parser.add_argument("--combined-json", default="selected_users_posts.json")
    parser.add_argument("--output-dir", default="weekly_docs")
    parser.add_argument("--delay", type=float, default=2.2)
    parser.add_argument("--timeout", type=int, default=35)
    parser.add_argument("--overlap", type=int, default=2)
    parser.add_argument("--page-limit", type=int, default=50)
    parser.add_argument("--skip-scrape", action="store_true", help="Only rebuild weekly docs from existing JSON.")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    start, end = week_bounds(args.week_of)
    label = week_key(start)
    combined_json = Path(args.combined_json)
    output_dir = Path(args.output_dir)

    if not args.skip_scrape:
        if not os.environ.get("NGA_COOKIE", "").strip():
            print("NGA_COOKIE is required unless --skip-scrape is used.")
            return 2
        run_step(
            [
                sys.executable,
                "update_latest_posts.py",
                "--combined-json",
                str(combined_json),
                "--delay",
                str(args.delay),
                "--timeout",
                str(args.timeout),
                "--overlap",
                str(args.overlap),
                "--page-limit",
                str(args.page_limit),
                "--allow-failures",
            ]
        )

    posts = load_posts(combined_json)
    week_posts = filter_week(posts, start, end)
    if not week_posts:
        print(f"No posts found for {label} ({start:%Y-%m-%d} to {end:%Y-%m-%d}).")
        return 1

    output_dir.mkdir(parents=True, exist_ok=True)
    week_doc = build_week_doc(label, start, end, week_posts, output_dir)
    week_md = write_week_markdown(label, start, end, week_posts, output_dir)

    market_cache = Path(f"market_cache_{start:%Y-%m}.json")
    if market_cache.exists():
        market_cache.unlink()

    analysis_doc = output_dir / f"{THREAD_TITLE}_发言逐条分析_{label}.docx"
    analysis_md = output_dir / f"{THREAD_TITLE}_发言逐条分析_{label}.md"
    run_step(
        [
            sys.executable,
            "build_june_analyzed_doc.py",
            "--input",
            str(combined_json),
            "--month",
            start.strftime("%Y-%m"),
            "--start-date",
            start.strftime("%Y-%m-%d"),
            "--end-date",
            end.strftime("%Y-%m-%d"),
            "--out",
            str(analysis_doc),
            "--markdown",
            str(analysis_md),
            "--market-cache",
            str(market_cache),
        ]
    )

    print(
        json.dumps(
            {
                "total_posts": len(posts),
                "week": label,
                "start": start.strftime("%Y-%m-%d"),
                "end": end.strftime("%Y-%m-%d"),
                "week_posts": len(week_posts),
                "latest_week": max(post.posted_at for post in week_posts),
                "week_doc": str(week_doc),
                "week_markdown": str(week_md),
                "analysis_doc": str(analysis_doc),
                "analysis_markdown": str(analysis_md),
            },
            ensure_ascii=False,
            indent=2,
        )
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
