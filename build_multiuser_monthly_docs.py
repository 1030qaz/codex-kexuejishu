#!/usr/bin/env python3
r"""
Scrape selected NGA authors from one thread and build monthly Word documents.

Usage:
  $env:NGA_COOKIE = "your logged-in NGA cookie"
  python .\build_multiuser_monthly_docs.py

The output is designed for later per-post analysis: each post is a compact
table with stable metadata and a clear body cell.
"""

from __future__ import annotations

import argparse
import json
import os
import re
import time
from collections import defaultdict
from dataclasses import asdict, dataclass
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from scrape_nga_user_thread import crawl


THREAD_ID = "45974302"
THREAD_TITLE = "科学技术打头阵"
THREAD_URL = f"https://bbs.nga.cn/read.php?tid={THREAD_ID}"

TARGET_USERS: list[tuple[str, str]] = [
    ("150058", "狼大"),
    ("67143809", "Rzzz二号机"),
    ("61395264", "村上吹树"),
    ("42162697", "包子music"),
    ("331181", "xbox"),
    ("65329649", "zippo578"),
    ("64648193", "一生一股绊倒铁盒"),
    ("60916468", "灰兔尾"),
    ("66908070", "夜骰2号机"),
    ("67152095", "小火机"),
    ("38090976", "sora496"),
    ("67177418", "小蝎子ICE二代"),
    ("5337595", "September_L"),
    ("67145714", "Plezl"),
]


@dataclass
class UserPost:
    uid: str
    username: str
    floor: str
    time: str
    posted_at: str
    body: str
    page: int
    url: str


def parse_time(value: str) -> datetime | None:
    match = re.search(r"(\d{4})[-/](\d{1,2})[-/](\d{1,2})\s+(\d{1,2}):(\d{1,2})", value)
    if not match:
        return None
    year, month, day, hour, minute = map(int, match.groups())
    try:
        return datetime(year, month, day, hour, minute)
    except ValueError:
        return None


def clean_body(value: str) -> str:
    text = value.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\[/?(?:quote|collapse|b|i|u|url|img|color|size|pid|uid)[^\]]*\]", "", text, flags=re.I)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()


def scrape_user(uid: str, username: str, cookie: str, delay: float, timeout: int, max_pages: int) -> list[UserPost]:
    url = f"{THREAD_URL}&authorid={uid}"
    raw_posts = crawl(url, cookie=cookie, delay=delay, timeout=timeout, max_pages=max_pages)
    posts: list[UserPost] = []
    for post in raw_posts:
        parsed = parse_time(post.time)
        if parsed is None:
            continue
        posts.append(
            UserPost(
                uid=uid,
                username=username,
                floor=post.floor,
                time=post.time,
                posted_at=parsed.isoformat(timespec="minutes"),
                body=clean_body(post.body),
                page=post.page,
                url=post.url,
            )
        )
    return posts


def load_or_scrape(
    output_json: Path,
    cookie: str,
    delay: float,
    timeout: int,
    max_pages: int,
    reuse: bool,
) -> list[UserPost]:
    if output_json.exists():
        data = json.loads(output_json.read_text(encoding="utf-8-sig"))
        existing_posts = [UserPost(**item) for item in data]
        if reuse:
            return existing_posts
        all_posts = sorted_posts(existing_posts)
    else:
        all_posts = []

    completed_uids = {post.uid for post in all_posts}
    for index, (uid, username) in enumerate(TARGET_USERS, start=1):
        if uid in completed_uids:
            print(f"[{index}/{len(TARGET_USERS)}] Skipping {username} ({uid}); already in {output_json}", flush=True)
            continue
        print(f"[{index}/{len(TARGET_USERS)}] Fetching {username} ({uid})", flush=True)
        user_posts = scrape_user(uid, username, cookie, delay, timeout, max_pages)
        print(f"  -> {len(user_posts)} posts", flush=True)
        all_posts.extend(user_posts)
        output_json.write_text(
            json.dumps([asdict(post) for post in sorted_posts(all_posts)], ensure_ascii=False, indent=2) + "\n",
            encoding="utf-8",
        )
        time.sleep(delay)

    return sorted_posts(all_posts)


def sorted_posts(posts: list[UserPost]) -> list[UserPost]:
    def key(post: UserPost) -> tuple[str, int, str]:
        floor_num = int(post.floor) if post.floor.isdigit() else 0
        return (post.posted_at, floor_num, post.uid)

    return sorted(posts, key=key)


def set_cell_shading(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def style_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Pt(50)
    section.bottom_margin = Pt(50)
    section.left_margin = Pt(54)
    section.right_margin = Pt(54)

    styles = document.styles
    for style_name in ["Normal", "Title", "Heading 1", "Heading 2"]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10)
    styles["Title"].font.size = Pt(20)
    styles["Heading 1"].font.size = Pt(15)


def add_post(document: Document, post: UserPost) -> None:
    is_wolf = post.uid == "150058"
    table = document.add_table(rows=2, cols=1)
    table.style = "Table Grid"

    header = table.rows[0].cells[0]
    set_cell_shading(header, "EAF3EA" if is_wolf else "F3F4F6")
    header_text = header.paragraphs[0]
    header_text.paragraph_format.space_after = Pt(0)
    run = header_text.add_run(f"{post.time}  |  {post.floor}楼  |  {post.username}  |  UID {post.uid}")
    run.bold = True
    run.font.size = Pt(11.5 if is_wolf else 9.5)
    run.font.color.rgb = RGBColor(22, 101, 52) if is_wolf else RGBColor(55, 65, 81)

    body_cell = table.rows[1].cells[0]
    body_paragraph = body_cell.paragraphs[0]
    body_paragraph.paragraph_format.line_spacing = 1.15
    body_paragraph.paragraph_format.space_after = Pt(0)
    for paragraph_index, paragraph_text in enumerate(post.body.split("\n\n")):
        paragraph = body_paragraph if paragraph_index == 0 else body_cell.add_paragraph()
        paragraph.paragraph_format.line_spacing = 1.15
        paragraph.paragraph_format.space_after = Pt(5)
        for line_index, line in enumerate(paragraph_text.split("\n")):
            if line_index:
                paragraph.add_run().add_break()
            body_run = paragraph.add_run(line.strip())
            body_run.font.size = Pt(12 if is_wolf else 10)
            if is_wolf:
                body_run.font.color.rgb = RGBColor(17, 24, 39)

    document.add_paragraph().paragraph_format.space_after = Pt(2)


def month_key(post: UserPost) -> str:
    return post.posted_at[:7]


def build_month_doc(month: str, posts: list[UserPost], output_dir: Path) -> Path:
    document = Document()
    style_document(document)

    title = document.add_paragraph()
    title.style = document.styles["Title"]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(f"{THREAD_TITLE}：{month} 发言整理")

    wolf_count = sum(1 for post in posts if post.uid == "150058")
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run(f"共 {len(posts)} 条；狼大 {wolf_count} 条；按发言时间顺序排列")
    sub_run.font.size = Pt(9.5)
    sub_run.font.color.rgb = RGBColor(107, 114, 128)

    document.add_heading("发言列表", level=1)
    for index, post in enumerate(posts, start=1):
        if index > 1 and (index - 1) % 80 == 0:
            document.add_section(WD_SECTION_START.NEW_PAGE)
        add_post(document, post)

    output_path = output_dir / f"{THREAD_TITLE}_发言整理_{month}.docx"
    document.save(output_path)
    return output_path


def build_monthly_docs(posts: list[UserPost], output_dir: Path) -> list[Path]:
    grouped: dict[str, list[UserPost]] = defaultdict(list)
    for post in sorted_posts(posts):
        grouped[month_key(post)].append(post)

    output_dir.mkdir(parents=True, exist_ok=True)
    paths: list[Path] = []
    for month in sorted(grouped):
        paths.append(build_month_doc(month, grouped[month], output_dir))
    return paths


def write_markdown(posts: list[UserPost], output_dir: Path) -> None:
    grouped: dict[str, list[UserPost]] = defaultdict(list)
    for post in sorted_posts(posts):
        grouped[month_key(post)].append(post)

    for month, month_posts in sorted(grouped.items()):
        lines = [f"# {THREAD_TITLE}：{month} 发言整理", ""]
        for post in month_posts:
            lines.append(f"## {post.time} | {post.floor}楼 | {post.username} | UID {post.uid}")
            lines.append("")
            lines.append(post.body)
            lines.append("")
        (output_dir / f"{THREAD_TITLE}_发言整理_{month}.md").write_text("\n".join(lines), encoding="utf-8")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build monthly docs for selected NGA users.")
    parser.add_argument("--combined-json", default="selected_users_posts.json")
    parser.add_argument("--output-dir", default="monthly_docs")
    parser.add_argument("--delay", type=float, default=0.5)
    parser.add_argument("--timeout", type=int, default=25)
    parser.add_argument("--max-pages", type=int, default=300)
    parser.add_argument("--reuse", action="store_true", help="Reuse existing combined JSON instead of scraping.")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    cookie = os.environ.get("NGA_COOKIE", "").strip()
    combined_json = Path(args.combined_json)
    output_dir = Path(args.output_dir)

    if not cookie and not args.reuse:
        print("NGA_COOKIE is required unless --reuse is used.")
        return 2

    posts = load_or_scrape(
        output_json=combined_json,
        cookie=cookie,
        delay=args.delay,
        timeout=args.timeout,
        max_pages=args.max_pages,
        reuse=args.reuse,
    )
    posts = sorted_posts(posts)
    combined_json.write_text(json.dumps([asdict(post) for post in posts], ensure_ascii=False, indent=2) + "\n", encoding="utf-8")

    doc_paths = build_monthly_docs(posts, output_dir)
    write_markdown(posts, output_dir)
    print(f"Saved {len(posts)} posts to {combined_json}")
    print(f"Generated {len(doc_paths)} Word documents in {output_dir}")
    for path in doc_paths:
        print(path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
