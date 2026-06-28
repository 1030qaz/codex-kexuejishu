#!/usr/bin/env python3
"""Distill Wolf's historical replies into a compact learning guide."""

from __future__ import annotations

import json
import re
from collections import Counter, defaultdict
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import load_workbook


TITLE = "狼大交易体系蒸馏"
SOURCE_PATTERN = "*20260626-1410*.xlsx"
OUTPUT_DIR = Path("分析参考")
DOCX_OUT = OUTPUT_DIR / "狼大交易体系蒸馏_20260626.docx"
MD_OUT = OUTPUT_DIR / "狼大交易体系蒸馏_20260626.md"
STATS_OUT = OUTPUT_DIR / "狼大交易体系蒸馏_20260626_stats.json"


CATEGORIES: dict[str, list[str]] = {
    "边界声明/身份定位": ["个人看法", "不做", "指导", "渠道", "账号", "讨论", "研究分享"],
    "市场环境/情绪温度": ["大盘", "指数", "下跌家数", "跌停", "情绪", "主力资金", "盘面", "安全系数", "认错"],
    "周期/节奏": ["周期", "调整浪", "交易日", "反弹波", "延伸浪", "阶段", "节奏", "时间"],
    "主线/产业逻辑": ["AI", "算力", "半导体", "芯片", "存储", "HBM", "PCB", "光", "白银", "有色", "医药", "赛道"],
    "买点/执行": ["挂单", "低吸", "到线", "不追", "追涨", "买", "抄", "接", "回踩"],
    "仓位/做T": ["仓位", "做T", "T", "加仓", "减仓", "定投", "锁仓", "回补", "撤单"],
    "技术结构": ["趋势", "支撑", "压力", "均线", "K", "上影", "下影", "揉搓线", "突破", "箱体"],
    "风控/反身性": ["风险", "止损", "割肉", "亏", "不碰", "问题不大", "利空", "泡沫", "兑现"],
    "散户教育/认知纠偏": ["看不懂", "纪律", "耐心", "情绪", "散户", "学习", "研究", "逻辑", "不要"],
}


@dataclass
class Reply:
    sheet: str
    posted_at: str
    text: str


def clean_text(value: Any) -> str:
    text = "" if value is None else str(value)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\[quote\].*?\[/quote\]", "", text, flags=re.I | re.S)
    text = re.sub(r"\[/?(?:b|i|u|url|img|color|size|pid|uid|tid|quote)[^\]]*\]", "", text, flags=re.I)
    text = re.sub(r"^Reply to .*?\)\s*", "", text, flags=re.I)
    text = re.sub(r"^Post by .*?\)\s*", "", text, flags=re.I)
    text = re.sub(r"\./mon_\d{6}/\d{2}/\S+\.(?:jpg|jpeg|png|gif|webp)", "", text, flags=re.I)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def parse_time(value: Any) -> str:
    if isinstance(value, datetime):
        return value.strftime("%Y-%m-%d %H:%M")
    text = "" if value is None else str(value).strip()
    match = re.search(r"(\d{4})[-/](\d{1,2})[-/](\d{1,2})\s+(\d{1,2}):(\d{1,2})", text)
    if not match:
        return text
    y, m, d, h, minute = map(int, match.groups())
    return f"{y:04d}-{m:02d}-{d:02d} {h:02d}:{minute:02d}"


def load_replies(path: Path) -> list[Reply]:
    wb = load_workbook(path, read_only=True, data_only=True)
    replies: list[Reply] = []
    seen: set[tuple[str, str]] = set()
    for ws in wb.worksheets:
        rows = ws.iter_rows(min_row=2, values_only=True)
        for row in rows:
            if not row or len(row) < 2:
                continue
            posted_at = parse_time(row[0])
            text = clean_text(row[1])
            if not posted_at or not text:
                continue
            key = (posted_at, text)
            if key in seen:
                continue
            seen.add(key)
            replies.append(Reply(ws.title, posted_at, text))
    replies.sort(key=lambda item: item.posted_at)
    return replies


def has_keyword(text: str, keyword: str) -> bool:
    lowered = text.lower()
    if keyword == "T":
        return bool(re.search(r"(做T|T进|T出|T掉|T回|T回来|日内T|T进去)", text))
    if keyword == "K":
        return bool(re.search(r"(K线|日K|周K|月K|\d{2,3}K)", text))
    if keyword == "光":
        return any(word in text for word in ["光模块", "光通信", "光芯片", "光华", "光伏"])
    return keyword.lower() in lowered


def score(text: str, keywords: list[str]) -> int:
    lowered = text.lower()
    return sum(1 for kw in keywords if has_keyword(text, kw))


def example_priority(reply: Reply, category_score: int) -> tuple[int, int, int]:
    year = int(reply.posted_at[:4]) if re.match(r"\d{4}", reply.posted_at) else 0
    recent_bonus = 3 if year >= 2025 else 1 if year >= 2020 else 0
    length_bonus = 2 if 35 <= len(reply.text) <= 220 else 0
    quote_penalty = -4 if "Reply to" in reply.text or "Post by" in reply.text else 0
    return category_score + recent_bonus + length_bonus + quote_penalty, year, len(reply.text)


def short_snippet(text: str, limit: int = 54) -> str:
    text = re.sub(r"\s+", " ", text).strip()
    if len(text) <= limit:
        return text
    return text[: limit - 1] + "…"


def analyze(replies: list[Reply]) -> dict[str, Any]:
    by_sheet = Counter(reply.sheet for reply in replies)
    by_year = Counter(reply.posted_at[:4] for reply in replies if re.match(r"\d{4}", reply.posted_at))
    category_counts: Counter[str] = Counter()
    category_candidates: dict[str, list[tuple[tuple[int, int, int], Reply]]] = defaultdict(list)
    keyword_counts: Counter[str] = Counter()

    for reply in replies:
        for category, keywords in CATEGORIES.items():
            s = score(reply.text, keywords)
            if s:
                category_counts[category] += 1
                category_candidates[category].append((example_priority(reply, s), reply))
        for keywords in CATEGORIES.values():
            for kw in keywords:
                if has_keyword(reply.text, kw):
                    keyword_counts[kw] += 1

    category_examples: dict[str, list[dict[str, str]]] = {}
    for category, candidates in category_candidates.items():
        selected = []
        used = set()
        for _, reply in sorted(
            candidates,
            key=lambda item: (item[0], item[1].posted_at, item[1].text[:40]),
            reverse=True,
        ):
            snippet = short_snippet(reply.text)
            if snippet in used or len(snippet) < 10:
                continue
            selected.append({"time": reply.posted_at, "sheet": reply.sheet, "snippet": snippet})
            used.add(snippet)
            if len(selected) >= 5:
                break
        category_examples[category] = selected

    return {
        "total": len(replies),
        "by_sheet": dict(by_sheet),
        "by_year": dict(sorted(by_year.items())),
        "category_counts": dict(category_counts.most_common()),
        "category_examples": category_examples,
        "keyword_counts": dict(keyword_counts.most_common(60)),
        "first_time": replies[0].posted_at if replies else "",
        "last_time": replies[-1].posted_at if replies else "",
    }


def lines_from_analysis(stats: dict[str, Any]) -> list[str]:
    total = stats["total"]
    start = stats["first_time"]
    end = stats["last_time"]
    top_categories = list(stats["category_counts"].items())[:6]
    top_keywords = "、".join(k for k, _ in list(stats["keyword_counts"].items())[:18])

    lines: list[str] = [
        f"# {TITLE}",
        "",
        f"> 数据源：`狼大回复汇总+20260626-1410&往期.xlsx`。清洗去重后共 {total} 条，时间跨度 {start} 至 {end}。",
        "> 本文只做个人学习和复盘框架整理，不构成投资建议，也不替代原作者本人判断。",
        "",
        "## 一句话蒸馏",
        "",
        "狼大的核心不是“报票”，而是一个以市场环境为底座、以主线产业逻辑为中枢、以位置和节奏为执行约束、以仓位和情绪纪律为安全阀的复盘系统。",
        "",
        "## 数据轮廓",
        "",
        "| 维度 | 内容 |",
        "| --- | --- |",
        f"| 覆盖记录 | {total} 条 |",
        f"| 时间跨度 | {start} 至 {end} |",
        f"| Sheet | {'；'.join(f'{k}: {v}' for k, v in stats['by_sheet'].items())} |",
        f"| 高频词 | {top_keywords} |",
        "",
        "## 狼大模型：五层判断链",
        "",
        "1. 先判市场温度：指数、涨跌家数、跌停数量、黄白线、期指升贴水、融资盘和资金情绪先于个股。",
        "2. 再判主线真假：不是当天涨得多就是主线，而是容量、逻辑、核心标的、风向标和可延展性同时成立。",
        "3. 再判阶段和筹码：放量看资金构成和筹码交换，缩量看一致性和下一步验证。",
        "4. 再判位置和节奏：调整周期、反弹波、回踩、支撑压力决定能不能动手。",
        "5. 再判执行方式：倾向挂单、低吸、到线买，不追涨，不在情绪最热时补票。",
        "6. 最后判复盘归因：赚少、没买到、买错、卖飞，都要回到原先条件和当时盘面，而不是事后脑补。",
        "",
        "## forum_post.txt 对模型的校正",
        "",
        "101楼整理帖不是普通发言流，而是系统性方法合集。它提示当前蒸馏需要从“日常发言总结”升级为“复盘表 + 做T条件 + 趋势生命周期 + 核心分层”的分析框架。",
        "",
        "- 市场环境要补多空评分、黄白线、期指升贴水、融资盘和龙虎榜等变量。",
        "- 执行动作要补做T适用条件、禁止条件和时间窗口，不能把做T写成默认动作。",
        "- 仓位管理要补被动止盈、利润垫、底仓、金字塔/倒金字塔和账户能力适配。",
        "- 主线判断要补核心、风向标、龙头、暗线和杂毛的分层。",
        "- 趋势分析要补量价筹码生命周期：埋伏、预期修正、承接换筹、爆发、大换手、逻辑变更。",
        "- 来源归属要更严格：整理帖内有图哥和其他用户内容，不能直接混成狼大本人观点。",
        "",
        "## 蒸馏出的心智模型",
        "",
        "### 多空温度表模型",
        "",
        "先判断市场能不能承接风险，再讨论板块和个股。指数、量能、下跌家数、跌停、黄白线、融资盘、期指升贴水和资金风格，是所有动作的底层开关。",
        "",
        "### 主线容量模型",
        "",
        "真正的主线不是涨幅最大，而是产业逻辑能延展、资金容量能容纳、核心标的能持续验证、补涨和分支能形成梯队。",
        "",
        "### 核心/风向标/龙头分层模型",
        "",
        "趋势行情里先分清板块核心、风向标和龙头。核心绑定板块生命力，风向标代表短期赚钱效应，龙头通常从风向标里产生。核心倒了，板块发散和杂毛补涨都要降级。",
        "",
        "### 量价筹码生命周期模型",
        "",
        "趋势龙头是结果，不是提前喊出来的。放量用来看资金构成和筹码交换，缩量用来看短期一致和下一步验证。",
        "",
        "### 条件触发模型",
        "",
        "观点必须绑定条件：到什么位置、出现什么量能、核心票怎么走、失效点在哪里。没有触发条件的发言，只能归为观察。",
        "",
        "### 仓位缓冲模型",
        "",
        "仓位不是表达信仰，而是对不确定性的缓冲。做T、回补、减仓、定投，本质是在承认自己不可能一次买在最低点。",
        "",
        "### 后验校正模型",
        "",
        "复盘不以盈亏本身定对错，而是检查原假设是否被盘面验证；错了要分清是方向错、时机错、仓位错、执行错，还是市场环境错。",
        "",
        "## 决策启发式",
        "",
        "- 如果盘面还混乱，先降低仓位预期，再寻找确定性。",
        "- 如果一个方向只有消息刺激、没有核心和容量，不把它当主线。",
        "- 如果核心不强，不从杂毛里强行寻找补涨确定性。",
        "- 如果要区分核心、风向标和龙头，先确认板块处于趋势行情而非长期横盘。",
        "- 如果上涨已经靠情绪推动，买点要后移到回踩或确认，不能追着情绪补票。",
        "- 如果个股逻辑成立但板块不配合，只能降低动作级别。",
        "- 如果买入理由来自“别人都在说”，先暂停，把产业逻辑、位置和失效条件写出来。",
        "- 如果盘中观点和盘后复盘冲突，以当时可见信息和原计划为准，不用事后结果倒推。",
        "- 如果做T三日内没有达到预期，要承认预期问题，不把做T仓位自动变成加仓。",
        "",
        "## 表达DNA与判断口吻",
        "",
        "- 常用纠偏式表达：先指出散户误区，再回到条件、位置和纪律。",
        "- 口吻偏交易现场：更关心能不能执行，而不是概念是否漂亮。",
        "- 经常用反问和否定句降低确定性幻觉，例如不追、不急、不做指导、看不懂先别动。",
        "- 强调时间戳：盘前、盘中、盘后不是同一类信息，不能混用。",
        "- 对技术分析的态度是工具化：K线、趋势、影线、量能必须放在市场环境里解释。",
        "",
        "## 诚实边界",
        "",
        "- 这份文档只能蒸馏公开回复中的判断框架，不能代表狼大本人，也不能预测其未来观点。",
        "- Excel 语料来自论坛回复，存在上下文缺失、图片缺失、被回复对象缺失的问题。",
        "- 本文短摘只用于证明某类表达存在，不做长篇原文搬运。",
        "- 涉及具体标的时，只能作为复盘学习材料，不构成荐股或买卖依据。",
        "",
        "## 九个能力模块",
        "",
    ]
    module_notes = {
        "边界声明/身份定位": "他反复强调只表达个人看法，不做带教或指导；这决定了学习时应抽框架，不应把单句当指令。",
        "市场环境/情绪温度": "大量判断先从下跌家数、跌停数、盘面清晰度、主力资金行为和恐慌扩散入手。",
        "周期/节奏": "常用调整浪、交易日数量、反弹波和阶段定位，判断当前更适合进攻、等待还是降低预期。",
        "主线/产业逻辑": "长期偏向能解释资金持续性的产业链逻辑，尤其关注AI、算力、半导体、存储、有色、医药等方向。",
        "买点/执行": "执行上偏机械化，强调挂单、到线、低吸、避免追涨，把冲动交易降到最低。",
        "仓位/做T": "通过仓位、做T、定投、回补等动作控制节奏，而不是一次性押方向。",
        "技术结构": "技术不是孤立图形，而是和趋势、量能、支撑压力、影线以及市场环境一起看。",
        "风控/反身性": "风险不只来自利空，也来自拥挤、兑现、情绪扩散、泡沫叙事和自己看不懂。",
        "散户教育/认知纠偏": "经常纠正散户把结果当逻辑、把消息当买点、把上涨当确定性的误区。",
    }
    for category, count in top_categories:
        lines.append(f"### {category}（命中 {count} 条）")
        lines.append("")
        lines.append(module_notes.get(category, "这是高频出现的行为或认知模块，适合作为后续分析标签。"))
        examples = stats["category_examples"].get(category, [])[:3]
        if examples:
            lines.append("")
            lines.append("短摘证据：")
            for item in examples:
                lines.append(f"- {item['time']}：{item['snippet']}")
        lines.append("")

    lines.extend(
        [
            "## 可复用交易规则",
            "",
            "### 市场环境",
            "",
            "- 当盘面还没清楚时，先降预期，不急着从单个票上找答案。",
            "- 下跌家数、跌停数、黄白线、量能和情绪扩散，是判断风险温度的优先变量。",
            "- 市场强时也要判断资金在做什么：是真主线进攻，还是权重护盘、题材脉冲或割肉情绪释放。",
            "",
            "### 主线与板块",
            "",
            "- 主线需要有产业逻辑、资金容量、核心标的、板块梯队和可持续叙事。",
            "- 先分清核心、风向标、龙头、暗线和杂毛；不同类型对应不同风险承受力。",
            "- 低位不天然安全，高位不天然危险，关键是逻辑是否兑现、资金是否继续、位置是否过热。",
            "- 不单纯讨论个股，个股必须放回板块、产业链和当日盘面里。",
            "",
            "### 买点与执行",
            "",
            "- 好票到线再挂，宁愿买不到，也不要因为看见上涨而追。",
            "- 先设条件，再等触发；不是先有情绪，再找理由。",
            "- 做T和加减仓是节奏工具，不是为了证明自己每天都要操作。",
            "- 做T需要指数箱体、黄白线、主线确认、量能和时间窗口配合；能力不足或无法盯盘时，优先底仓与被动止盈。",
            "",
            "### 风控与心态",
            "",
            "- 赚少是交易成本的一部分；为了赚满而破坏纪律，才是长期亏损根源。",
            "- 看不懂主力资金在干什么时，不要把噪音翻译成确定性。",
            "- 每一次失败都要区分：方向错、时机错、仓位错、执行错，还是市场环境错配。",
            "- 被动止盈优先保护利润垫，卖飞不是主要风险，破坏纪律才是主要风险。",
            "",
            "## 狼大式每日复盘模板",
            "",
            "1. 今日市场温度：指数、量能、涨跌家数、跌停数、黄白线、期指和情绪是否扩散。",
            "2. 今日主线判断：资金真正进攻的方向是什么，核心、风向标、龙头是否互相验证。",
            "3. 今日阶段定位：处于埋伏、预期修正、换筹、爆发、大换手还是退潮。",
            "4. 今日操作条件：哪些方向只观察，哪些方向到线才做，哪些方向不能追。",
            "5. 今日风险点：拥挤、兑现、假突破、低位陷阱、消息刺激后衰减。",
            "6. 次日验证：看量能、核心标的、板块联动和原假设失效条件。",
            "",
            "## 反向禁区",
            "",
            "- 不要把单条回复当成买卖指令。",
            "- 不要脱离日期、时段和当时盘面复用观点。",
            "- 不要只看关键词，不看条件、位置、仓位和失效。",
            "- 不要把“逻辑好”直接等同于“马上会涨”。",
            "- 不要把图哥、整理者或其他用户的内容直接归为狼大观点。",
            "- 不要在长期横盘板块里硬找趋势核心、风向标和龙头。",
            "- 不要为了追求完整复刻而忽略自己的资金量、承受力和执行能力。",
            "",
            "## 后续接入项目的方式",
            "",
            "1. 给每条发言增加四个字段：市场温度、主线判断、动作条件、失效条件。",
            "2. 给每日总结增加两个字段：T+5后验观察、T+20后验观察。",
            "3. 把高频模块作为标签：市场环境、主线、周期、执行、仓位、技术、风险、认知纠偏。",
            "4. 分析文档里只展开有条件、有对象、有验证路径的发言；闲聊和情绪表达保留原文即可。",
        ]
    )
    return lines


def set_font(run: Any, size: float, bold: bool = False, color: str | None = None) -> None:
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    styles = doc.styles
    for name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.25
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.color.rgb = RGBColor.from_string("2E74B5")
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.color.rgb = RGBColor.from_string("2E74B5")
    styles["Heading 3"].font.size = Pt(12)
    styles["Heading 3"].font.color.rgb = RGBColor.from_string("1F4D78")


def add_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    widths = [Inches(1.45), Inches(5.05)]
    for idx, text in enumerate(["模块", "内容"]):
        cell = table.rows[0].cells[idx]
        cell.width = widths[idx]
        shade_cell(cell, "E8EEF5")
        run = cell.paragraphs[0].add_run(text)
        set_font(run, 10, bold=True, color="0B2545")
    for label, detail in rows:
        row = table.add_row().cells
        row[0].width = widths[0]
        row[1].width = widths[1]
        shade_cell(row[0], "F2F4F7")
        set_font(row[0].paragraphs[0].add_run(label), 9.5, bold=True)
        set_font(row[1].paragraphs[0].add_run(detail), 9.5)
    doc.add_paragraph()


def build_doc(lines: list[str], stats: dict[str, Any]) -> None:
    doc = Document()
    style_doc(doc)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(TITLE)
    set_font(run, 20, bold=True, color="0B2545")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(
        subtitle.add_run(f"基于 {stats['total']} 条历史回复；{stats['first_time']} 至 {stats['last_time']}"),
        9.5,
        color="6B7280",
    )

    data_rows = [
        ("数据源", "狼大回复汇总+20260626-1410&往期.xlsx"),
        ("覆盖范围", f"{stats['first_time']} 至 {stats['last_time']}"),
        ("清洗后记录", f"{stats['total']} 条"),
        ("用途边界", "只做个人学习和复盘框架，不构成投资建议。"),
    ]
    add_table(doc, data_rows)

    for line in lines:
        if not line or line.startswith("# ") or line.startswith("> "):
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            set_font(p.add_run(line[2:]), 10.5)
        elif re.match(r"\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            set_font(p.add_run(re.sub(r"^\d+\. ", "", line)), 10.5)
        elif line.startswith("| "):
            continue
        else:
            p = doc.add_paragraph()
            set_font(p.add_run(line), 10.5)
    doc.save(DOCX_OUT)


def main() -> int:
    candidates = [p for p in Path(".").glob(SOURCE_PATTERN)]
    if not candidates:
        raise FileNotFoundError(f"No workbook matching {SOURCE_PATTERN}")
    source = candidates[0]
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    replies = load_replies(source)
    stats = analyze(replies)
    lines = lines_from_analysis(stats)
    MD_OUT.write_text("\n".join(lines) + "\n", encoding="utf-8")
    STATS_OUT.write_text(json.dumps(stats, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    build_doc(lines, stats)
    print(json.dumps({"source": source.name, "total": stats["total"], "docx": str(DOCX_OUT), "markdown": str(MD_OUT)}, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
