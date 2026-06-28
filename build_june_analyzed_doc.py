#!/usr/bin/env python3
r"""
Build the June analyzed document for the NGA thread "科学技术打头阵".

Output design:
- keep all June posts in chronological order;
- analyze all A/B class posts with the same framework;
- highlight wolf's posts and make their analysis fuller;
- remove repetitive "小白学习笔记";
- resolve stock codes and common abbreviations to full names and directions.
"""

from __future__ import annotations

import argparse
import calendar
import html
import json
import re
import time
import urllib.parse
from collections import Counter, defaultdict
from dataclasses import dataclass
from datetime import datetime, timedelta
from pathlib import Path
from typing import Any

import requests
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

requests.packages.urllib3.disable_warnings(requests.packages.urllib3.exceptions.InsecureRequestWarning)


THREAD_ID = "45974302"
THREAD_TITLE = "科学技术打头阵"
WOLF_UID = "150058"
DAY_SUMMARY_REFERENCE = Path("分析参考/分析揉搓线及日内总结.txt")

INDEXES = {
    "1.000001": "上证指数",
    "0.399001": "深证成指",
    "0.399006": "创业板指",
    "1.000688": "科创50",
    "1.000852": "中证1000",
}

TENCENT_INDEXES = {
    "sh000001": "上证指数",
    "sz399001": "深证成指",
    "sz399006": "创业板指",
    "sh000688": "科创50",
    "sh000852": "中证1000",
}

TRADING_KEYWORDS = [
    "大盘", "指数", "上证", "深成", "创业板", "科创", "中证", "黄线", "白线",
    "板块", "主线", "方向", "细分", "科技", "光", "算力", "液冷", "AI",
    "半导体", "芯片", "封装", "PCB", "有色", "金属", "白酒", "券商",
    "银行", "医药", "消费", "机器人", "军工", "港股", "北向",
    "个股", "龙头", "中军", "补涨", "杂毛", "涨停", "跌停",
    "量能", "成交", "放量", "缩量", "换手", "资金", "承接", "兑现",
    "低吸", "追高", "高位", "低位", "仓", "减仓", "加仓", "止损",
    "止盈", "T", "做T", "回补", "风险", "分歧", "退潮", "轮动",
    "预期", "超预期", "不及预期", "利好", "利空", "落地",
    "K", "红K", "黑K", "揉搓线", "上影", "下影", "缺口", "回踩",
    "突破", "破位", "均线", "支撑", "压力", "V", "岛形",
]

SECTOR_KEYWORDS = {
    "核心科技/国产算力": ["科技", "核心科技", "国产算力", "华为", "昇腾", "算力", "服务器", "PCB"],
    "光通信/光模块": ["光模块", "光通信", "CPO", "光上游", "光芯片", "光器件", "光迅", "中际", "新易盛"],
    "半导体设备/材料/封测": ["半导体", "芯片", "封装", "封测", "设备", "材料", "长电", "通富", "北方华创", "中微"],
    "存储/HBM": ["存储", "HBM", "DRAM", "NAND", "德明利", "佰维", "江波龙", "兆易创新"],
    "PCB上游/载板材料": ["电子布", "铜箔", "树脂", "ABF", "BT载板", "玻纤", "覆铜板", "CCL"],
    "AI应用/软件": ["AI应用", "AI 软", "软件", "大模型", "传媒", "游戏"],
    "液冷/散热": ["液冷", "散热", "英维克", "强瑞"],
    "电源/燃气轮机/算电协同": ["电源", "燃气轮机", "算电协同", "固态变压器", "AIDC", "数据中心供电"],
    "有色金属": ["有色", "金属", "铜", "铝", "小金属", "贵金属"],
    "消费/白酒": ["消费", "白酒", "食品", "旅游"],
    "金融权重": ["券商", "银行", "保险", "金融", "权重"],
    "医药": ["医药", "药", "创新药", "医疗"],
    "机器人": ["机器人", "减速器", "电机"],
}

ACTION_KEYWORDS = {
    "观察": ["观察", "看", "等", "确认", "不急"],
    "低吸": ["低吸", "低挂", "回落", "接"],
    "持有": ["持有", "锁仓", "拿", "坚定"],
    "减仓": ["减仓", "降仓", "跑", "卖", "出清", "止盈"],
    "止损": ["止损", "破位"],
    "做T": ["做T", "T", "回补"],
    "不操作": ["不做", "别追", "不追", "不要追", "严禁"],
}

LAYER_KEYWORDS = {
    "核心": ["核心", "中军", "定价权", "核心标的"],
    "风向标": ["风向标", "领涨", "赚钱效应", "先起来"],
    "龙头": ["龙头", "龙一", "最强", "趋势领跑"],
    "暗线": ["暗线", "预期差", "隐藏预期", "细分延伸"],
    "杂毛": ["杂毛", "蹭", "低位杂毛", "非核心"],
}

STAGE_KEYWORDS = {
    "埋伏": ["埋伏", "平铺", "试错", "底仓"],
    "预期修正": ["预期", "上修", "业绩", "中报", "超预期", "不及预期"],
    "承接换筹": ["承接", "换筹", "交换筹码", "获利盘", "接盘", "换手"],
    "爆发": ["爆发", "加速", "突破", "放量拉", "主升"],
    "大换手": ["大换手", "放天量", "天量", "复杂顶", "滞涨"],
    "退潮": ["退潮", "兑现", "走弱", "破位", "跌破", "出清"],
}

ACCOUNT_KEYWORDS = {
    "适合躺底仓/被动止盈": ["躺", "底仓", "被动止盈", "不能盯盘", "长期", "BOLL"],
    "适合日内做T": ["做T", "日内T", "T出", "T回", "高抛低吸"],
    "需要降低仓位": ["减仓", "降仓", "仓位", "利润垫", "风险"],
    "能力边界提醒": ["看你自己的能力", "能力", "不要乱动", "别追", "不追"],
}

TERM_DEFS = {
    "主线": "持续获得资金选择、容量足够、逻辑可延展的方向，不等同于当天涨得多。",
    "黄线": "分时图里更接近中小盘/平均股表现的线，可辅助判断题材股强弱。",
    "白线": "按权重计算的指数线，可辅助判断大票和权重护盘强弱。",
    "量能": "成交量和成交额，代表资金参与程度；没有量能的形态可靠性较弱。",
    "低吸": "在回落、支撑或分歧处买入，而不是上涨后情绪化追高。",
    "追高": "价格已经明显上涨后再买，常见风险是买在短线兑现点。",
    "分歧": "资金意见不一致，常表现为冲高回落、放量震荡、板块内部强弱分化。",
    "预期差": "实际消息或盘面相对市场原先预期的偏离，关键在位置和资金是否提前埋伏。",
    "揉搓线": "上下影线明显、实体偏小的 K 线结构，表示多空拉扯，必须结合趋势和量能。",
    "止损": "触发预设风险条件后退出，避免亏损扩大。",
    "止盈": "达到目标或跌破保护位后锁定利润。",
    "回踩": "突破后价格回落测试支撑或关键位置。",
    "突破": "价格或指数上穿关键压力位，需要量能和板块核心跟随确认。",
}

STOCK_ALIAS_DIRECTION = {
    "英子": ("英维克", "002837", "液冷/数据中心温控"),
    "英维克": ("英维克", "002837", "液冷/数据中心温控"),
    "寒武": ("寒武纪", "688256", "国产 AI 芯片/算力"),
    "寒武纪": ("寒武纪", "688256", "国产 AI 芯片/算力"),
    "海光": ("海光信息", "688041", "国产 CPU/GPU/算力芯片"),
    "中芯": ("中芯国际", "688981", "晶圆制造/半导体制造"),
    "长电": ("长电科技", "600584", "先进封装/封测"),
    "通富": ("通富微电", "002156", "先进封装/封测"),
    "华天": ("华天科技", "002185", "封装测试"),
    "甬矽": ("甬矽电子", "688362", "封装测试"),
    "北方": ("北方华创", "002371", "半导体设备"),
    "北方华创": ("北方华创", "002371", "半导体设备"),
    "中微": ("中微公司", "688012", "半导体设备"),
    "拓荆": ("拓荆科技", "688072", "半导体设备"),
    "华海": ("华海清科", "688120", "半导体设备/CMP"),
    "鼎龙": ("鼎龙股份", "300054", "半导体材料/CMP 抛光垫"),
    "华丰": ("华丰科技", "688629", "高速连接器/算力硬件"),
    "高新": ("高新发展", "000628", "华为昇腾/算力集成"),
    "拓维": ("拓维信息", "002261", "华为昇腾/软件与算力"),
    "神码": ("神州数码", "000034", "华为鲲鹏/昇腾生态"),
    "神州数码": ("神州数码", "000034", "华为鲲鹏/昇腾生态"),
    "华大": ("华大九天", "301269", "EDA/芯片设计软件"),
    "概伦": ("概伦电子", "688206", "EDA/芯片设计软件"),
    "华工": ("华工科技", "000988", "光通信/激光设备"),
    "光迅": ("光迅科技", "002281", "光通信/光模块"),
    "中际": ("中际旭创", "300308", "光模块/CPO"),
    "大哥": ("中际旭创", "300308", "光模块/CPO"),
    "赵姨": ("兆易创新", "603986", "存储/MCU/半导体设计"),
    "姨": ("兆易创新", "603986", "存储/MCU/半导体设计"),
    "zycx": ("兆易创新", "603986", "存储/MCU/半导体设计"),
    "新易盛": ("新易盛", "300502", "光模块/CPO"),
    "二哥": ("新易盛", "300502", "光模块/CPO"),
    "xys": ("新易盛", "300502", "光模块/CPO"),
    "龙蟠": ("龙蟠科技", "603906", "材料/电池材料"),
    "强瑞": ("强瑞技术", "301128", "液冷/测试设备/算力硬件"),
    "工富": ("工业富联", "601138", "AI 服务器/算力硬件"),
    "工业富联": ("工业富联", "601138", "AI 服务器/算力硬件"),
    "沪电": ("沪电股份", "002463", "PCB/服务器高速板"),
    "护垫": ("沪电股份", "002463", "PCB/服务器高速板"),
    "hdgf": ("沪电股份", "002463", "PCB/服务器高速板"),
    "兴森": ("兴森科技", "002436", "PCB/IC载板"),
    "沃尔": ("沃尔核材", "002130", "线缆/高速连接/材料"),
    "胜宏": ("胜宏科技", "300476", "PCB/AI服务器板"),
    "shkj": ("胜宏科技", "300476", "PCB/AI服务器板"),
    "德明利": ("德明利", "001309", "存储模组/存储芯片"),
    "dml": ("德明利", "001309", "存储模组/存储芯片"),
    "澜起": ("澜起科技", "688008", "内存接口芯片/算力互连"),
    "蓝旗": ("澜起科技", "688008", "内存接口芯片/算力互连"),
    "lqkj": ("澜起科技", "688008", "内存接口芯片/算力互连"),
    "中科飞测": ("中科飞测", "688361", "半导体检测设备"),
    "KFC": ("中科飞测", "688361", "半导体检测设备"),
    "华峰测控": ("华峰测控", "688200", "半导体测试设备"),
    "盛美上海": ("盛美上海", "688082", "半导体清洗设备"),
    "沪硅产业": ("沪硅产业", "688126", "半导体硅片材料"),
    "立昂微": ("立昂微", "605358", "半导体硅片/功率器件"),
    "中兴": ("中兴通讯", "000063", "通信设备/算力网络"),
    "zte": ("中兴通讯", "000063", "通信设备/算力网络"),
    "英伟达": ("英伟达", "NVDA", "美股 AI GPU/算力龙头"),
}

IMAGE_BLACK_TALK = {
    "铁盒": "半导体设备方向，图片中用于指代半导体设备/相关 ETF 语境。",
    "宝贝王国": "bbwg，图片中的论坛黑话简称。",
    "dml": "德明利，偏存储模组/存储芯片方向。",
    "姨": "兆易创新，偏存储/MCU/半导体设计方向。",
    "赵姨": "兆易创新，偏存储/MCU/半导体设计方向。",
    "大哥": "中际旭创，光模块/CPO 方向。",
    "二哥": "新易盛，光模块/CPO 方向。",
    "牢弟": "图片中指向 PCB 估值龙头语境，需结合原文判断是否指胜宏科技。",
    "护垫": "沪电股份，PCB/服务器高速板方向。",
    "沪龟": "沪硅产业，半导体大硅片材料方向。",
    "荆轲": "拓荆科技，半导体设备方向。",
    "KFC": "中科飞测，半导体检测设备方向。",
    "华莱士": "华峰测控，半导体测试设备方向。",
    "小美": "盛美上海，半导体清洗设备方向。",
}

COMPUTE_CHAIN_RULES = [
    ("基础骨架/血管", ["电子布", "铜箔", "树脂", "PCB", "沪电", "胜宏", "深南", "生益"], "对应算力硬件的物理载体，逻辑是电子布/铜箔/树脂压制成 PCB，再承载 GPU、存储和光模块。"),
    ("算力大脑", ["GPU", "AI芯片", "算力芯片", "寒武纪", "海光", "昇腾", "HBM", "存储"], "对应半导体、GPU/AI 芯片、HBM/存储与先进封装，重点看算力供给、国产替代和封装瓶颈。"),
    ("高速互连", ["光模块", "CPO", "光芯片", "光纤", "光互连", "交换机"], "对应芯片间高速通信，图片框架把光模块/光纤网络/CPO 视为厂内高速公路。"),
    ("整机与基础设施", ["AI服务器", "服务器", "液冷", "电源", "燃气轮机", "算电协同"], "对应 AI 服务器、液冷、电源和算电协同，重点看从硬件堆叠到能源支撑能否闭环。"),
    ("商业兑现", ["算力租赁", "AIDC", "数据中心", "算力中心"], "对应算力中心建设、出租和下游 AI 公司使用，分析时要区分硬件涨价、建设成本和租赁收益。"),
]

SCARCITY_KEYWORDS = ["紧缺", "缺货", "涨价", "供需", "排产", "扩产", "瓶颈", "国产替代"]
SCARCITY_AREAS = "存储、PCB上游/ABF/BT载板、光纤、被动元件、CPU、光上游、算力租赁、国产GPU、国产交换机芯片"
RICE_LAW_KEYWORDS = ["稻定律", "华为", "TSV", "Chiplet", "先进封装", "CPO", "HBM", "系统总线", "光互连", "3D堆叠"]


@dataclass
class MarketDay:
    date: str
    indexes: dict[str, dict[str, float | str]]


class StockResolver:
    def __init__(self, cache_path: Path) -> None:
        self.cache_path = cache_path
        self.session = requests.Session()
        self.session.trust_env = False
        if cache_path.exists():
            self.cache: dict[str, Any] = json.loads(cache_path.read_text(encoding="utf-8-sig"))
        else:
            self.cache = {}

    def save(self) -> None:
        self.cache_path.write_text(json.dumps(self.cache, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")

    def direction_for_name(self, name: str) -> str:
        for _, (full, _, direction) in STOCK_ALIAS_DIRECTION.items():
            if name == full:
                return direction
        if any(k in name for k in ["光", "旭创", "易盛"]):
            return "光通信/光模块"
        if any(k in name for k in ["封", "电", "微"]):
            return "半导体/电子"
        if any(k in name for k in ["算", "联", "信息"]):
            return "算力/软件信息"
        return "待人工确认所属方向"

    def search(self, query: str) -> dict[str, str] | None:
        if query in self.cache:
            return self.cache[query]
        if query in STOCK_ALIAS_DIRECTION:
            full, code, direction = STOCK_ALIAS_DIRECTION[query]
            payload = {"alias": query, "name": full, "code": code, "direction": direction, "source": "内置映射"}
            self.cache[query] = payload
            return payload

        url = (
            "http://searchapi.eastmoney.com/api/suggest/get?"
            f"input={urllib.parse.quote(query)}&type=14&token=1"
        )
        try:
            response = self.session.get(
                url,
                timeout=10,
                headers={"User-Agent": "Mozilla/5.0", "Referer": "http://quote.eastmoney.com/"},
            )
            response.raise_for_status()
            data = response.json().get("QuotationCodeTable", {}).get("Data") or []
        except Exception:
            data = []
        result = None
        for item in data:
            if item.get("SecurityTypeName") in {"沪A", "深A", "科创板", "创业板", "美股"}:
                name = item.get("Name", "")
                code = item.get("Code", "")
                result = {
                    "alias": query,
                    "name": name,
                    "code": code,
                    "direction": self.direction_for_name(name),
                    "source": "东方财富搜索",
                }
                break
        self.cache[query] = result
        return result

    def detect(self, body: str) -> list[dict[str, str]]:
        candidates: set[str] = set()
        for alias in STOCK_ALIAS_DIRECTION:
            if alias in body:
                candidates.add(alias)
        for code in re.findall(r"(?<!\d)((?:60|68|00|30)\d{4})(?!\d)", body):
            candidates.add(code)
        found = []
        for query in sorted(candidates, key=len, reverse=True):
            item = self.search(query)
            if item:
                found.append(item)
        deduped = []
        seen = set()
        for item in found:
            key = (item["code"], item["name"])
            if key not in seen:
                seen.add(key)
                deduped.append(item)
        return deduped[:8]


def clean_text(value: str) -> str:
    text = html.unescape(value or "")
    text = re.sub(r"\[/?(?:quote|collapse|b|i|u|url|img|color|size|pid|uid)[^\]]*\]", "", text, flags=re.I)
    text = re.sub(r"Reply(?: to)?(?: Post by)?", "Reply", text)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r" *\n *", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def parse_time(value: str) -> datetime:
    return datetime.strptime(value[:16], "%Y-%m-%d %H:%M")


def time_stage(dt: datetime) -> str:
    hhmm = dt.hour * 100 + dt.minute
    if hhmm < 930:
        return "盘前/开盘前"
    if hhmm <= 1000:
        return "开盘确认段"
    if hhmm < 1130:
        return "早盘"
    if hhmm < 1300:
        return "午间"
    if hhmm < 1430:
        return "午后"
    if hhmm <= 1500:
        return "尾盘"
    return "收盘后"


def load_posts(path: Path, month: str, start_date: str | None = None, end_date: str | None = None) -> list[dict[str, Any]]:
    posts = json.loads(path.read_text(encoding="utf-8-sig"))
    result = []
    for post in posts:
        posted_at = post.get("posted_at", "")
        if month != "all" and not posted_at.startswith(month):
            continue
        day = posted_at[:10]
        if start_date and day < start_date:
            continue
        if end_date and day > end_date:
            continue
        item = dict(post)
        item["body"] = clean_text(item.get("body", ""))
        item["_dt"] = parse_time(item["time"])
        result.append(item)
    result.sort(key=lambda p: (p["_dt"], int(p["floor"]) if str(p.get("floor", "")).isdigit() else 0, p["uid"]))
    return result


def _date_dash(value: str) -> str:
    return f"{value[:4]}-{value[4:6]}-{value[6:8]}"


def fetch_tencent_klines(begin: str, end: str) -> dict[str, MarketDay]:
    begin_dash = _date_dash(begin)
    end_dash = _date_dash(end)
    session = requests.Session()
    session.trust_env = False
    by_day: dict[str, MarketDay] = {}

    for code, name in TENCENT_INDEXES.items():
        url = (
            "https://web.ifzq.gtimg.cn/appstock/app/fqkline/get?"
            f"param={code},day,{begin_dash},{end_dash},320,qfq"
        )
        response = None
        for attempt in range(1, 5):
            try:
                response = session.get(url, timeout=20, headers={"User-Agent": "Mozilla/5.0"})
                response.raise_for_status()
                break
            except Exception as exc:
                if attempt == 4:
                    print(f"Warning: failed fetching Tencent {name}: {exc}")
                else:
                    time.sleep(1.2 * attempt)
        if response is None:
            continue

        payload = response.json().get("data", {}).get(code, {})
        previous_close: float | None = None
        for line in payload.get("day", []):
            if len(line) < 6:
                continue
            date, open_, close, high, low, volume = line[:6]
            if not (begin_dash <= date <= end_dash):
                continue
            open_value = float(open_)
            close_value = float(close)
            high_value = float(high)
            low_value = float(low)
            volume_value = float(volume)
            if previous_close:
                change = close_value - previous_close
                pct = change / previous_close * 100
                amplitude = (high_value - low_value) / previous_close * 100
            else:
                change = 0.0
                pct = 0.0
                amplitude = 0.0
            by_day.setdefault(date, MarketDay(date=date, indexes={}))
            by_day[date].indexes[name] = {
                "open": open_value,
                "close": close_value,
                "high": high_value,
                "low": low_value,
                "volume": volume_value,
                "amount": 0.0,
                "amplitude": amplitude,
                "pct_change": pct,
                "change": change,
                "turnover": 0.0,
                "source": "腾讯财经日线",
            }
            previous_close = close_value

        quote = payload.get("qt", {}).get(code)
        if quote and len(quote) > 37:
            raw_date = quote[30][:8]
            date = _date_dash(raw_date)
            if begin_dash <= date <= end_dash:
                close_value = float(quote[3])
                previous_close = float(quote[4])
                open_value = float(quote[5])
                high_value = float(quote[33])
                low_value = float(quote[34])
                volume_value = float(quote[36])
                amount_value = float(quote[37]) * 10000
                by_day.setdefault(date, MarketDay(date=date, indexes={}))
                by_day[date].indexes[name] = {
                    "open": open_value,
                    "close": close_value,
                    "high": high_value,
                    "low": low_value,
                    "volume": volume_value,
                    "amount": amount_value,
                    "amplitude": (high_value - low_value) / previous_close * 100 if previous_close else 0.0,
                    "pct_change": float(quote[32]),
                    "change": float(quote[31]),
                    "turnover": 0.0,
                    "source": "腾讯财经实时行情",
                }

    return by_day


def fetch_eastmoney_klines(begin: str, end: str) -> dict[str, MarketDay]:
    session = requests.Session()
    session.trust_env = False
    by_day: dict[str, MarketDay] = {}
    failures: list[str] = []

    for secid, name in INDEXES.items():
        url = (
            "https://push2his.eastmoney.com/api/qt/stock/kline/get"
            f"?secid={secid}&fields1=f1,f2,f3,f4,f5,f6"
            "&fields2=f51,f52,f53,f54,f55,f56,f57,f58,f59,f60,f61"
            f"&klt=101&fqt=1&beg={begin}&end={end}&rtntype=6"
        )
        response = None
        last_error: Exception | None = None
        for attempt in range(1, 5):
            try:
                response = session.get(
                    url,
                    timeout=25,
                    headers={"User-Agent": "Mozilla/5.0", "Referer": "https://quote.eastmoney.com/"},
                    verify=False,
                )
                response.raise_for_status()
                break
            except Exception as exc:
                last_error = exc
                time.sleep(1.5 * attempt)
        if response is None:
            failures.append(f"{name}: {last_error}")
            continue
        data = response.json().get("data") or {}
        for line in data.get("klines", []):
            parts = line.split(",")
            if len(parts) < 11:
                continue
            date, open_, close, high, low, volume, amount, amplitude, pct, change, turnover = parts[:11]
            by_day.setdefault(date, MarketDay(date=date, indexes={}))
            by_day[date].indexes[name] = {
                "open": float(open_),
                "close": float(close),
                "high": float(high),
                "low": float(low),
                "volume": float(volume),
                "amount": float(amount),
                "amplitude": float(amplitude),
                "pct_change": float(pct),
                "change": float(change),
                "turnover": float(turnover),
                "source": "东方财富日线",
            }

    if failures:
        print("Warning: Eastmoney index fetch failed for " + "; ".join(failures[:3]))
    return by_day


def _next_day_yyyymmdd(day: str) -> str:
    return (datetime.strptime(day, "%Y-%m-%d") + timedelta(days=1)).strftime("%Y%m%d")


def _merge_market_days(base: dict[str, MarketDay], fresh: dict[str, MarketDay]) -> dict[str, MarketDay]:
    for day, md in fresh.items():
        base.setdefault(day, MarketDay(date=day, indexes={}))
        base[day].indexes.update(md.indexes)
    return base


def _has_non_eastmoney_sources(days: dict[str, MarketDay]) -> bool:
    for md in days.values():
        for item in md.indexes.values():
            if item.get("source") != "东方财富日线":
                return True
    return False


def fetch_index_klines(begin: str, end: str, cache_path: Path) -> dict[str, MarketDay]:
    cached: dict[str, MarketDay] = {}
    if cache_path.exists():
        payload = json.loads(cache_path.read_text(encoding="utf-8-sig"))
        cached = {day: MarketDay(date=day, indexes=item["indexes"]) for day, item in payload.items()}

    effective_end = min(end, datetime.now().strftime("%Y%m%d"))
    if begin > effective_end:
        return cached

    fetch_begin = begin
    if cached:
        if _has_non_eastmoney_sources(cached):
            fresh_full = fetch_eastmoney_klines(begin, effective_end)
            if fresh_full:
                cached = _merge_market_days(cached, fresh_full)
                cache_path.write_text(
                    json.dumps({day: {"date": md.date, "indexes": md.indexes} for day, md in cached.items()}, ensure_ascii=False, indent=2) + "\n",
                    encoding="utf-8",
                )
        latest_cached = max(cached)
        end_dash = _date_dash(effective_end)
        if latest_cached >= end_dash and not _has_non_eastmoney_sources(cached):
            return cached
        if latest_cached >= end_dash and _has_non_eastmoney_sources(cached):
            fresh_full = fetch_eastmoney_klines(begin, effective_end)
            if fresh_full:
                by_day = _merge_market_days(cached, fresh_full)
                cache_path.write_text(
                    json.dumps({day: {"date": md.date, "indexes": md.indexes} for day, md in by_day.items()}, ensure_ascii=False, indent=2) + "\n",
                    encoding="utf-8",
                )
                return by_day
            print("Warning: Eastmoney refresh unavailable; keeping cached fallback market data.")
            return cached
        fetch_begin = _next_day_yyyymmdd(latest_cached)

    fresh = fetch_eastmoney_klines(fetch_begin, effective_end)
    if fresh:
        by_day = _merge_market_days(cached, fresh)
    elif cached:
        fallback = fetch_tencent_klines(fetch_begin, effective_end)
        if fallback:
            print("Warning: Eastmoney index data unavailable; falling back to Tencent Finance for missing market days only.")
            by_day = _merge_market_days(cached, fallback)
        else:
            print("Info: no additional Eastmoney trading-day data returned for missing range; keeping cached market data.")
            by_day = cached
    else:
        by_day = fetch_tencent_klines(begin, end)
        if not by_day:
            print("Info: no index kline data returned; this may be a non-trading holiday/weekend range.")
        else:
            print("Warning: Eastmoney index data unavailable; falling back to Tencent Finance.")

    cache_path.write_text(
        json.dumps({day: {"date": md.date, "indexes": md.indexes} for day, md in by_day.items()}, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    return by_day


def market_summary(day: str, market_days: dict[str, MarketDay]) -> str:
    md = market_days.get(day)
    if not md:
        return "未取得当日指数数据。"
    parts = []
    for name in ["上证指数", "深证成指", "创业板指", "科创50", "中证1000"]:
        item = md.indexes.get(name)
        if item:
            parts.append(f"{name}{item['pct_change']:+.2f}%")
    up_count = sum(1 for item in md.indexes.values() if float(item["pct_change"]) > 0)
    amount = md.indexes.get("上证指数", {}).get("amount")
    volume_note = f"；上证成交额约{float(amount)/1e8:.0f}亿元" if amount else ""
    style_note = "指数整体偏强" if up_count >= 3 else "指数分化或偏弱"
    return "，".join(parts) + volume_note + f"；{style_note}。"


def parse_referenced_dates(body: str, post_dt: datetime) -> list[str]:
    dates: set[str] = set()
    for match in re.finditer(r"(20\d{2})[-/年](\d{1,2})[-/月](\d{1,2})日?", body):
        year, month, day = map(int, match.groups())
        try:
            dates.add(datetime(year, month, day).strftime("%Y-%m-%d"))
        except ValueError:
            pass
    for match in re.finditer(r"(?<!\d)(\d{1,2})月(\d{1,2})日", body):
        month, day = map(int, match.groups())
        try:
            dates.add(datetime(post_dt.year, month, day).strftime("%Y-%m-%d"))
        except ValueError:
            pass
    current_day = post_dt.strftime("%Y-%m-%d")
    return sorted(day for day in dates if day != current_day)


def market_days_for_month(month: str, current_month: str, current_market_days: dict[str, MarketDay]) -> dict[str, MarketDay]:
    if month == current_month:
        return current_market_days
    cache_path = Path(f"market_cache_{month}.json")
    year, month_num = map(int, month.split("-"))
    last_day = calendar.monthrange(year, month_num)[1]
    begin = f"{year}{month_num:02d}01"
    end = f"{year}{month_num:02d}{last_day:02d}"
    return fetch_index_klines(begin, end, cache_path)


def referenced_market_notes(body: str, post_dt: datetime, current_market_days: dict[str, MarketDay]) -> list[str]:
    notes: list[str] = []
    current_month = post_dt.strftime("%Y-%m")
    for day in parse_referenced_dates(body, post_dt):
        month = day[:7]
        month_days = market_days_for_month(month, current_month, current_market_days)
        notes.append(f"{day}：{market_summary(day, month_days)}")
    return notes


def keyword_hits(body: str) -> list[str]:
    hits: list[str] = []
    lowered = body.lower()
    for keyword in TRADING_KEYWORDS:
        if keyword == "T":
            if re.search(r"(?<![A-Za-z])(?:做T|正T|反T|T出|T回|T仓|日内T)(?![A-Za-z])", body):
                hits.append(keyword)
            continue
        if keyword == "K":
            if re.search(r"(?:K线|红K|黑K|日K|周K|月K)", body):
                hits.append(keyword)
            continue
        if keyword.lower() in lowered:
            hits.append(keyword)
    return hits


def has_substantive_signal(body: str) -> bool:
    if detect_sectors(body) or technical_notes(body) or reference_notes(body):
        return True
    action = detect_action(body)
    if action != "观察/未明确":
        return True
    if re.search(r"(低开|高开|放量|缩量|支撑|压力|缺口|仓位|减仓|加仓|止盈|止损|买|卖|突破|回踩|黄线|白线)", body):
        return True
    return False


def classify_post(post: dict[str, Any]) -> tuple[str, str]:
    body = post["body"]
    if not body:
        return "C", "空正文或无有效信息。"
    hits = keyword_hits(body)
    substantive = has_substantive_signal(body)
    if post["uid"] == WOLF_UID and len(hits) >= 2:
        return "A", "狼大发言且包含多个交易/盘面关键词，适合重点详析。"
    if post["uid"] == WOLF_UID and substantive:
        return "B", "狼大发言，交易信息较弱或偏互动，但仍保留结构化分析。"
    if post["uid"] == WOLF_UID:
        return "C", "狼大发言但缺少可还原的交易条件，保留原文不展开。"
    if len(hits) >= 2:
        return "A", "包含多个交易/盘面关键词，适合完整分析。"
    if hits and substantive:
        return "B", "包含少量交易关键词，适合作为上下文分析。"
    return "C", "偏互动、闲聊或情绪表达，不展开。"


def detect_terms(body: str) -> list[tuple[str, str]]:
    return [(term, desc) for term, desc in TERM_DEFS.items() if term in body][:8]


def detect_sectors(body: str) -> list[str]:
    return [sector for sector, keywords in SECTOR_KEYWORDS.items() if any(keyword in body for keyword in keywords)]


def detect_action(body: str) -> str:
    for action, keywords in ACTION_KEYWORDS.items():
        if any(keyword in body for keyword in keywords):
            return action
    return "观察/未明确"


def detect_layer(body: str) -> str:
    hits = [name for name, keywords in LAYER_KEYWORDS.items() if any(keyword in body for keyword in keywords)]
    return "、".join(hits[:3]) if hits else "不足以判断"


def detect_trend_stage(body: str) -> str:
    hits = [name for name, keywords in STAGE_KEYWORDS.items() if any(keyword in body for keyword in keywords)]
    return "、".join(hits[:3]) if hits else "不足以判断"


def detect_account_fit(body: str) -> str:
    hits = [name for name, keywords in ACCOUNT_KEYWORDS.items() if any(keyword in body for keyword in keywords)]
    return "、".join(hits[:3]) if hits else "不足以判断"


def t_trade_notes(body: str) -> list[str]:
    notes: list[str] = []
    has_t = re.search(r"做T|日内T|T出|T回|T仓|高抛低吸", body)
    if not has_t:
        return notes
    if "黄线" in body or "白线" in body:
        notes.append("做T条件：需要先确认黄白线强弱，黄线强更偏题材/中小票，白线强则日内T成功率要降级。")
    if "放量" in body or "缩量" in body:
        notes.append("做T条件：分时上涨放量、回调缩量更健康；放量杀跌或黄白线交织时要谨慎。")
    if "支撑" in body or "压力" in body or "缺口" in body:
        notes.append("做T条件：优先贴近支撑、缺口和压力位做计划，不能把开盘冲动当成买卖点。")
    notes.append("仓位纪律：做T仓位不能自动变成加仓；三日内不达预期要承认假设问题。")
    return notes[:4]


def technical_notes(body: str) -> list[str]:
    notes = []
    if "揉搓线" in body or "上影" in body or "下影" in body:
        notes.append("涉及揉搓线/影线结构：必须先判断当前是上涨趋势还是下跌趋势，再结合红黑 K、上下影线长短和量能。")
    if "放量" in body or "缩量" in body:
        notes.append("涉及量能：放量说明资金参与提高，缩量说明分歧或参与不足，要看放量后能否延续。")
    if "缺口" in body:
        notes.append("涉及缺口：需要观察是否回补，以及是否形成跳空岛形等结构风险。")
    if "突破" in body or "回踩" in body:
        notes.append("涉及突破/回踩：关键是量能、核心标的和板块跟随是否确认。")
    if "黄线" in body or "白线" in body:
        notes.append("涉及黄白线：黄线强偏题材和中小票，白线强偏权重和大票，不能混用。")
    if any(word in body for word in ["升水", "贴水", "期指", "A50", "沪深300", "中证1000"]):
        notes.append("涉及期指/升贴水：只能作为盘前和开盘前后情绪参考，仍需盘中量能、黄白线和板块核心验证。")
    return notes


def reference_notes(body: str) -> list[str]:
    notes: list[str] = []
    for stage, keywords, note in COMPUTE_CHAIN_RULES:
        if any(keyword.lower() in body.lower() for keyword in keywords):
            notes.append(f"{stage}：{note}")
    if any(keyword in body for keyword in SCARCITY_KEYWORDS) and any(
        keyword in body for keyword in ["AI", "算力", "存储", "PCB", "光", "芯片", "服务器"]
    ):
        notes.append(f"紧缺环节：图片将 {SCARCITY_AREAS} 列为 AI 产业链相对紧缺方向；分析时要看供需缺口、业绩兑现和股价位置是否已提前定价。")
    slang_hits = []
    for alias, meaning in IMAGE_BLACK_TALK.items():
        if alias in body:
            slang_hits.append(f"{alias}={meaning}")
    if slang_hits:
        notes.append("黑话/简称：" + "；".join(slang_hits[:6]))
    if any(keyword in body for keyword in RICE_LAW_KEYWORDS):
        notes.append(
            "华为稻定律框架：重点不是单一材料或单一设备，而是从材料、EDA/仿真、制造封装、互连、CPO/HBM 到整机系统的延迟压缩和系统级协同。"
        )
    return notes[:6]


def plain_language(post: dict[str, Any], cls: str, stocks: list[dict[str, str]]) -> str:
    if cls == "C":
        return "这条不作为交易知识重点，保留原文即可。"
    sectors = detect_sectors(post["body"])
    action = detect_action(post["body"])
    actor = "狼大" if post["uid"] == WOLF_UID else post["username"]
    if sectors or stocks:
        text = f"{actor}这条主要围绕"
    else:
        text = f"{actor}这条主要是在补充盘面结构或交易纪律。"
    if sectors:
        text += f"{'、'.join(sectors)}"
    if stocks:
        stock_text = "；".join(f"{s['alias']}→{s['name']}({s['code']})/{s['direction']}" for s in stocks)
        text += f"；涉及标的：{stock_text}"
    if sectors or stocks:
        text += "。"
    if action != "观察/未明确":
        text += f" 动作倾向：{action}。"
    return text


def make_analysis(post: dict[str, Any], market_days: dict[str, MarketDay], resolver: StockResolver) -> dict[str, Any]:
    day = post["posted_at"][:10]
    cls, reason = classify_post(post)
    stocks = resolver.detect(post["body"]) if cls != "C" else []
    market = market_summary(day, market_days) if cls != "C" else ""
    referenced_markets = referenced_market_notes(post["body"], post["_dt"], market_days) if cls != "C" else []
    sectors = detect_sectors(post["body"])
    layer = detect_layer(post["body"]) if cls != "C" else "不展开"
    trend_stage = detect_trend_stage(post["body"]) if cls != "C" else "不展开"
    account_fit = detect_account_fit(post["body"]) if cls != "C" else "不展开"
    terms = detect_terms(post["body"])
    tech = technical_notes(post["body"])
    t_notes = t_trade_notes(post["body"]) if cls != "C" else []
    refs = reference_notes(post["body"])
    action = detect_action(post["body"]) if cls != "C" else "不展开"
    chain = []
    if cls != "C":
        chain = [
            f"时间位置：{time_stage(post['_dt'])}，盘前、盘中和收盘后观点不能混用。",
            f"层级判断：{'、'.join(sectors) if sectors else '未明显点名板块，侧重交易纪律/盘面结构'}。",
            f"狼大分层：{layer}；只有趋势行情里核心/风向标/龙头分层才有较高解释力。",
            f"阶段定位：{trend_stage}；需用放量/缩量、筹码交换和核心标的验证。",
        ]
        if stocks:
            chain.append(f"个股/简称：{'；'.join(f'{s['alias']}→{s['name']}({s['code']})/{s['direction']}' for s in stocks)}。")
        if action != "观察/未明确":
            chain.append(f"动作条件：更接近“{action}”，需要回到原文触发条件和当时盘面。")
        if account_fit != "不足以判断":
            chain.append(f"账户适配：{account_fit}，不能脱离盯盘能力和资金承受力复刻。")
    return {
        "class": cls,
        "reason": reason,
        "translation": plain_language(post, cls, stocks),
        "market": market,
        "referenced_markets": referenced_markets,
        "sectors": sectors,
        "layer": layer,
        "trend_stage": trend_stage,
        "account_fit": account_fit,
        "stocks": stocks,
        "terms": terms,
        "tech": tech,
        "t_notes": t_notes,
        "refs": refs,
        "action": action,
        "chain": chain,
        "hypothesis": hypothesis_card(post, {
            "sectors": sectors,
            "stocks": stocks,
            "tech": tech,
            "action": action,
            "layer": layer,
            "trend_stage": trend_stage,
            "referenced_markets": referenced_markets,
        }) if cls != "C" else "",
        "postmortem": postmortem_plan({
            "sectors": sectors,
            "stocks": stocks,
            "tech": tech,
            "layer": layer,
            "trend_stage": trend_stage,
        }) if cls != "C" else "",
        "memory": memory_tag(post, {
            "sectors": sectors,
            "action": action,
            "tech": tech,
            "layer": layer,
            "trend_stage": trend_stage,
        }) if cls != "C" else "",
        "quality": quality_guardrail({
            "action": action,
            "tech": tech,
            "layer": layer,
            "trend_stage": trend_stage,
        }) if cls != "C" else "",
    }


def set_cell_shading(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_font(run: Any, size: float, bold: bool = False, color: str | None = None) -> None:
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Pt(48)
    section.bottom_margin = Pt(48)
    section.left_margin = Pt(54)
    section.right_margin = Pt(54)
    styles = doc.styles
    for name in ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10)
    styles["Title"].font.size = Pt(20)
    styles["Heading 1"].font.size = Pt(15)
    styles["Heading 2"].font.size = Pt(12)


def add_cell_line(cell: Any, head: str, value: str, head_color: str = "111827") -> None:
    if not value:
        return
    paragraph = cell.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2)
    label = paragraph.add_run(f"{head}：")
    set_font(label, 9.5, bold=True, color=head_color)
    body = paragraph.add_run(value)
    set_font(body, 9.5, color="374151")


def day_summary_reference_note() -> str:
    if DAY_SUMMARY_REFERENCE.exists():
        return "参考《分析揉搓线及日内总结.txt》、论坛上下文及交易复盘框架；仅做学习复盘，不作荐股结论。"
    return "参考日内总结与交易复盘框架；仅做学习复盘，不作荐股结论。"


def framework_note() -> str:
    return (
        "框架吸收 signal-postmortem、trader-memory-core、trade-hypothesis-ideator、"
        "technical-analyst、sector-analyst、market-environment-analysis、data-quality-checker 的方法论："
        "先提炼可证伪假设，再记录验证/失效条件，最后用市场环境、板块轮动、技术结构和数据质检闭环复盘。"
        "本版额外接入 wolf-perspective：多空温度表、核心/风向标/龙头分层、量价筹码生命周期、做T条件和账户适配。"
    )


def top_items(counter: Counter[str], limit: int = 5) -> list[str]:
    return [name for name, count in counter.most_common(limit) if count > 0]


def summarize_time_windows(day_posts: list[dict[str, Any]]) -> str:
    windows = {
        "盘前/开盘确认": 0,
        "早盘": 0,
        "午间/午后": 0,
        "尾盘/收盘后": 0,
    }
    for post in day_posts:
        stage = time_stage(post["_dt"])
        if stage in windows:
            windows[stage] += 1
        elif "开盘" in stage:
            windows["盘前/开盘确认"] += 1
        elif "早盘" in stage:
            windows["早盘"] += 1
        elif "午" in stage:
            windows["午间/午后"] += 1
        else:
            windows["尾盘/收盘后"] += 1
    return "；".join(f"{name}{count}条" for name, count in windows.items() if count)


def compact_label(post: dict[str, Any], analysis: dict[str, Any]) -> str:
    sectors = "、".join(analysis["sectors"][:2])
    stocks = "、".join(f"{s['alias']}→{s['name']}" for s in analysis["stocks"][:2])
    label = sectors or stocks or analysis["translation"][:36]
    return f"{post['time'][11:16]} {post['username']}：{label}"


def evidence_quality(analysis: dict[str, Any]) -> str:
    score = 0
    if analysis["sectors"]:
        score += 1
    if analysis["stocks"]:
        score += 1
    if analysis["tech"]:
        score += 1
    if analysis["action"] != "观察/未明确":
        score += 1
    if analysis["referenced_markets"]:
        score += 1
    if score >= 4:
        return "较强：方向、动作、技术或市场环境有多项交叉证据。"
    if score >= 2:
        return "中等：有可分析线索，但仍需盘面和后续走势验证。"
    return "偏弱：更多是观点片段或情绪反馈，不能单独形成交易结论。"


def hypothesis_card(post: dict[str, Any], analysis: dict[str, Any]) -> str:
    sectors = "、".join(analysis["sectors"][:3]) or "未明确点名板块"
    action = analysis["action"] if analysis["action"] != "观察/未明确" else "观察/等待确认"
    stage = time_stage(post["_dt"])
    layer = analysis.get("layer") or "不足以判断"
    trend_stage = analysis.get("trend_stage") or "不足以判断"
    hypothesis = f"若{sectors}在{stage}后仍有量能、核心标的和板块联动，且分层定位({layer})与阶段({trend_stage})被盘面确认，则该方向更可能延续；否则只当作当时盘面观察。"
    validation = "验证：看次日/后续5个交易日是否继续放量、核心/风向标/龙头是否主动、补涨是否有承接。"
    invalidation = "失效：缩量反抽、核心转弱、只剩非核心脉冲，做T条件未满足，或原文触发条件没有出现。"
    return f"{hypothesis} 动作倾向：{action}。{validation} {invalidation} 证据质量：{evidence_quality(analysis)}"


def postmortem_plan(analysis: dict[str, Any]) -> str:
    if not (analysis["sectors"] or analysis["stocks"] or analysis["tech"]):
        return ""
    objects = analysis["sectors"][:2] + [f"{s['name']}({s['code']})" for s in analysis["stocks"][:2]]
    target = "、".join(objects) if objects else "该观点涉及方向"
    return (
        f"把{target}记录为观察样本，复盘T+5/T+20表现；"
        "若方向兑现则归因于量能/主线/技术结构/核心分层，若失败则标注假突破、情绪脉冲、核心走弱或市场环境错配。"
    )


def memory_tag(post: dict[str, Any], analysis: dict[str, Any]) -> str:
    tags = [post["username"], time_stage(post["_dt"])]
    if analysis["sectors"]:
        tags.extend(analysis["sectors"][:2])
    if analysis["action"] != "观察/未明确":
        tags.append(analysis["action"])
    if analysis["tech"]:
        tags.append("技术结构")
    if analysis.get("layer") and analysis["layer"] != "不足以判断":
        tags.append("分层:" + analysis["layer"])
    if analysis.get("trend_stage") and analysis["trend_stage"] != "不足以判断":
        tags.append("阶段:" + analysis["trend_stage"])
    return " / ".join(tags) + "；记录原文、当时市场环境、验证条件、失效条件和后续复盘结论。"


def quality_guardrail(analysis: dict[str, Any]) -> str:
    checks = ["日期/时段不能混用", "指数和板块结论需对应同一交易日", "简称必须补全后再学习"]
    if analysis["action"] == "观察/未明确":
        checks.append("未出现明确动作时不得改写成买卖建议")
    if not analysis["tech"]:
        checks.append("缺少技术结构时不得强行解释K线")
    if analysis.get("layer") == "不足以判断":
        checks.append("未识别核心/风向标/龙头时不得强行分层")
    if analysis.get("trend_stage") == "不足以判断":
        checks.append("未识别趋势阶段时只做观察，不套生命周期")
    return "；".join(checks) + "。"


def market_regime_label(market: str, valuable_count: int, c_count: int) -> str:
    if "偏强" in market:
        regime = "偏风险偏好"
    elif "偏弱" in market:
        regime = "偏防御/风险收缩"
    else:
        regime = "震荡观察"
    confidence = "中等" if valuable_count >= c_count else "偏低"
    return f"环境标签：{regime}；结论置信度{confidence}。"


def sector_rotation_read(sectors: list[str], sector_counter: Counter[str]) -> str:
    if not sectors:
        return "板块轮动：有效板块线索不足，不能强行判断周期阶段。"
    leader = "、".join(sectors[:3])
    breadth = "较集中" if sector_counter[sectors[0]] >= 3 else "偏分散"
    return f"板块轮动：相对强线索集中在{leader}；发言宽度{breadth}，需比较主线核心、中军与补涨梯队是否同步。"


def technical_scenario_read(techs: list[str], sectors: list[str]) -> str:
    base = "基准情景：趋势和量能继续配合，相关方向延续观察。"
    bull = "强化情景：放量突破或回踩不破，核心标的带动板块扩散。"
    bear = "失效情景：缩量、长上影、跌破支撑或核心标的走弱。"
    if not techs:
        return "技术情景：当日可用技术线索不足；先看趋势、支撑/压力、量能和影线是否在后续发言中被确认。"
    focus = "、".join(sectors[:2]) if sectors else "相关方向"
    return f"技术情景（{focus}）：{base} {bull} {bear}"


def daily_hypotheses(valuable: list[tuple[dict[str, Any], dict[str, Any]]], limit: int = 3) -> list[str]:
    cards = []
    for post, analysis in valuable:
        if analysis["sectors"] or analysis["stocks"] or analysis["tech"] or analysis["action"] != "观察/未明确":
            cards.append(compact_label(post, analysis) + "；" + hypothesis_card(post, analysis))
        if len(cards) >= limit:
            break
    return cards


def categorize_adjustments(valuable: list[tuple[dict[str, Any], dict[str, Any]]]) -> dict[str, list[str]]:
    buckets: dict[str, list[str]] = {"看多/持有": [], "观望": [], "减仓/出清": [], "风险标的": []}
    for post, analysis in valuable:
        label = compact_label(post, analysis)
        action = analysis["action"]
        text = post["body"]
        if action in {"持有", "低吸", "做T"} or any(word in text for word in ["锁仓", "继续", "拿住", "低吸", "回补"]):
            buckets["看多/持有"].append(label)
        elif action in {"减仓", "止损", "不操作"} or any(word in text for word in ["减仓", "出清", "止损", "别追", "不碰", "风险"]):
            if action == "不操作" or "别追" in text or "不碰" in text:
                buckets["风险标的"].append(label)
            else:
                buckets["减仓/出清"].append(label)
        else:
            buckets["观望"].append(label)
    for key in buckets:
        buckets[key] = buckets[key][:5]
    return buckets


def format_adjustment_buckets(buckets: dict[str, list[str]]) -> str:
    parts = []
    for name in ["看多/持有", "观望", "减仓/出清", "风险标的"]:
        items = buckets.get(name) or []
        parts.append(f"{name}：" + ("；".join(items) if items else "不足以判断"))
    return "\n".join(parts)


def reusable_rules(sectors: list[str], techs: list[str], terms: list[str], valuable_count: int) -> str:
    rules = [
        "判断主线真假：优先看量能是否随核心方向放大、核心标的是否主动走强、板块梯队是否延续。",
        "判断板块持续性：看核心与中军是否共振，补涨是否有承接，非主线脉冲是否只是在搅动市场。",
        "判断个股能不能追：先看位置、量能和上方压力；上涨趋势里的长上影更适合等低接，不适合情绪追高。",
        "判断低位方向是不是陷阱：看是否有持续资金、业绩或产业逻辑，而不是只看低位和当天涨幅。",
    ]
    if sectors:
        rules.append(f"次日观察：继续跟踪{'、'.join(sectors[:3])}的量能、核心标的和板块联动。")
    elif valuable_count == 0:
        rules.append("次日观察：当日有效信息不足，先回到指数、量能、黄白线和主线核心确认。")
    if techs or "揉搓线" in terms:
        rules.append("技术框架：揉搓线、影线和缺口必须结合趋势位置与量能，不孤立解释。")
    return " ".join(rules)


def build_day_summary(day: str, day_posts: list[dict[str, Any]], market_days: dict[str, MarketDay], resolver: StockResolver) -> dict[str, str]:
    analyses = [(post, make_analysis(post, market_days, resolver)) for post in day_posts]
    valuable = [(post, analysis) for post, analysis in analyses if analysis["class"] != "C"]
    wolf_valuable = [(post, analysis) for post, analysis in valuable if post["uid"] == WOLF_UID]
    sector_counter: Counter[str] = Counter()
    action_counter: Counter[str] = Counter()
    layer_counter: Counter[str] = Counter()
    stage_counter: Counter[str] = Counter()
    term_counter: Counter[str] = Counter()
    tech_counter: Counter[str] = Counter()
    stock_counter: Counter[str] = Counter()
    for _, analysis in valuable:
        sector_counter.update(analysis["sectors"])
        if analysis["action"] != "观察/未明确":
            action_counter.update([analysis["action"]])
        if analysis["layer"] != "不足以判断":
            layer_counter.update(analysis["layer"].split("、"))
        if analysis["trend_stage"] != "不足以判断":
            stage_counter.update(analysis["trend_stage"].split("、"))
        term_counter.update(term for term, _ in analysis["terms"])
        tech_counter.update(analysis["tech"])
        stock_counter.update(f"{s['alias']}→{s['name']}" for s in analysis["stocks"])

    sectors = top_items(sector_counter)
    actions = top_items(action_counter, 4)
    layers = top_items(layer_counter, 4)
    stages = top_items(stage_counter, 4)
    terms = top_items(term_counter, 4)
    techs = top_items(tech_counter, 3)
    stocks = top_items(stock_counter, 6)
    market = market_summary(day, market_days)

    if sectors:
        tone = f"{market} {market_regime_label(market, len(valuable), len(day_posts) - len(valuable))} 有效讨论集中在{'、'.join(sectors[:4])}，主线判断需回到量能、核心标的、风向标/龙头和板块联动验证。"
    else:
        tone = f"{market} {market_regime_label(market, len(valuable), len(day_posts) - len(valuable))} 有效交易信息偏少，更多是互动、情绪反馈或既有观点补充，不足以强行归纳主线。"
    tone += f" 当日A/B类有效发言{len(valuable)}条，C类原文保留{len(day_posts) - len(valuable)}条。"

    adjustment = format_adjustment_buckets(categorize_adjustments(valuable))
    adjustment += "\n" + sector_rotation_read(sectors, sector_counter)
    if layers:
        adjustment += "\n狼大分层线索：" + "、".join(layers)
    if stages:
        adjustment += "\n趋势阶段线索：" + "、".join(stages)

    core_sources = wolf_valuable[:3] or valuable[:3]
    if core_sources:
        core = " -> ".join(compact_label(post, analysis) for post, analysis in core_sources)
        core = f"发言链条：{core}。重点看哪些资金在进攻、哪些方向只是搅动市场，以及量能是否验证主线真假。"
    else:
        core = "不足以判断；当日没有足够可串联的 A/B 类发言，不强行脑补资金进攻和对手盘关系。"

    if techs:
        technical = "；".join(techs[:4]) + " " + technical_scenario_read(techs, sectors)
    else:
        technical = technical_scenario_read(techs, sectors)
    if layers or stages:
        technical += f" 狼大框架校正：分层={ '、'.join(layers) if layers else '不足以判断' }；阶段={ '、'.join(stages) if stages else '不足以判断' }。"

    risks = []
    if techs:
        risks.append("技术结构：" + "；".join(techs[:2]))
    if terms:
        risks.append("高频术语：" + "、".join(terms))
    if not risks and len(valuable) <= 3:
        risks.append("有效交易信息偏少，不能从闲聊和情绪反馈里推导操作。")
    if len(day_posts) - len(valuable) > len(valuable):
        risks.append("C类内容占比较高，阅读时应避免把论坛情绪当作盘面结论。")
    risk = " ".join(risks) if risks else "主要风险在于把单条观点脱离触发条件使用；仍需结合当日指数、量能和板块核心表现验证。"

    dynamics = (
        f"时段分布：{summarize_time_windows(day_posts)}。"
        f" 当日发言总数{len(day_posts)}条，其中狼大发言{sum(1 for p in day_posts if p['uid'] == WOLF_UID)}条。"
        " 阅读顺序建议：先看盘前/盘后预判，再看盘中验证和情绪变化，最后看是否形成闭环判断。"
        " 复盘闭环：把明确观点纳入T+5/T+20观察，区分真兑现、假信号、错过机会和市场环境错配。"
    )

    if terms or stocks:
        glossary_items = []
        if terms:
            glossary_items.extend(terms[:5])
        if stocks:
            glossary_items.extend(stocks[:5])
        glossary = "；".join(glossary_items)
    else:
        glossary = "当日未提取到高频特殊术语、缩写或黑话；不足以判断。"

    hypotheses = daily_hypotheses(valuable)
    if hypotheses:
        core += "\n可证伪假设：\n" + "\n".join(f"{idx}. {item}" for idx, item in enumerate(hypotheses, start=1))
    else:
        core += "\n可证伪假设：不足以判断。"

    rules = reusable_rules(sectors, techs, terms, len(valuable))
    rules += " 交易记忆：保留观点来源、原始条件、验证日、失效条件和复盘结论。 数据质检：日期、简称、指数环境、板块归因和动作词必须逐项核对。"
    rules += " 狼大框架：先看多空温度，再看核心/风向标/龙头分层，再看量价筹码阶段；做T必须满足黄白线、量能、支撑压力和账户能力条件。"

    return {
        "市场定调": tone,
        "板块/个股调仓建议": adjustment,
        "核心博弈逻辑": core,
        "技术与盘口语言": technical,
        "风险/机会提示": risk,
        "时段论坛动态综合总结": dynamics,
        "特殊术语与黑话释义": glossary,
        "可复用交易框架": rules,
    }


def add_day_summary(doc: Document, day: str, day_posts: list[dict[str, Any]], market_days: dict[str, MarketDay], resolver: StockResolver) -> None:
    doc.add_heading("每日操盘内参", level=3)
    note = doc.add_paragraph()
    r = note.add_run(day_summary_reference_note())
    set_font(r, 9, color="6B7280")
    summary = build_day_summary(day, day_posts, market_days, resolver)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.columns[0].width = Pt(96)
    table.columns[1].width = Pt(390)
    header = table.rows[0].cells
    for idx, text in enumerate(["模块", "总结"]):
        set_cell_shading(header[idx], "E5E7EB")
        run = header[idx].paragraphs[0].add_run(text)
        set_font(run, 9.5, bold=True, color="111827")
    for index, (head, value) in enumerate(summary.items(), start=1):
        row = table.add_row().cells
        set_cell_shading(row[0], "F3F4F6")
        label = row[0].paragraphs[0].add_run(f"{index}. {head}")
        set_font(label, 9.2, bold=True, color="111827")
        for part_idx, part in enumerate(value.split("\n")):
            para = row[1].paragraphs[0] if part_idx == 0 else row[1].add_paragraph()
            para.paragraph_format.space_after = Pt(2)
            body = para.add_run(part)
            set_font(body, 9.2, color="374151")
    doc.add_paragraph()


def add_post_card(doc: Document, post: dict[str, Any], analysis: dict[str, Any]) -> None:
    is_wolf = post["uid"] == WOLF_UID
    has_analysis = analysis["class"] != "C"
    table = doc.add_table(rows=3 if has_analysis else 2, cols=1)
    table.style = "Table Grid"
    header = table.rows[0].cells[0]
    set_cell_shading(header, "DFF3E4" if is_wolf else "F3F4F6")
    hp = header.paragraphs[0]
    hp.paragraph_format.space_after = Pt(0)
    run = hp.add_run(f"{post['username']}  |  {post['time']}  |  {post['floor']}楼  |  UID {post['uid']}  |  {analysis['class']}类")
    set_font(run, 11.5 if is_wolf else 9.5, bold=True, color="166534" if is_wolf else "374151")

    body_cell = table.rows[1].cells[0]
    for idx, paragraph_text in enumerate(post["body"].split("\n\n")):
        para = body_cell.paragraphs[0] if idx == 0 else body_cell.add_paragraph()
        para.paragraph_format.line_spacing = 1.15
        para.paragraph_format.space_after = Pt(4)
        for line_idx, line in enumerate(paragraph_text.split("\n")):
            if line_idx:
                para.add_run().add_break()
            rr = para.add_run(line.strip())
            set_font(rr, 11 if is_wolf else 9.5)

    if has_analysis:
        analysis_cell = table.rows[2].cells[0]
        set_cell_shading(analysis_cell, "F7FCF8" if is_wolf else "FAFAFA")
        title = "狼大发言详析" if is_wolf else "发言分析"
        p = analysis_cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(title)
        set_font(r, 10.5 if is_wolf else 9.8, bold=True, color="166534" if is_wolf else "111827")
        add_cell_line(analysis_cell, "一句话还原", analysis["translation"])
        if analysis["sectors"]:
            add_cell_line(analysis_cell, "板块/细分方向", "、".join(analysis["sectors"]))
        if analysis["stocks"]:
            add_cell_line(
                analysis_cell,
                "个股简称补全",
                "；".join(f"{s['alias']} → {s['name']}（{s['code']}，{s['direction']}，{s['source']}）" for s in analysis["stocks"]),
            )
        if analysis["referenced_markets"]:
            add_cell_line(analysis_cell, "跨日市场环境", "；".join(analysis["referenced_markets"]))
        if analysis["terms"]:
            add_cell_line(analysis_cell, "术语解释", "；".join(f"{term}：{desc}" for term, desc in analysis["terms"]))
        if analysis["tech"]:
            add_cell_line(analysis_cell, "技术/量能结构", "；".join(analysis["tech"]))
        if analysis["layer"] != "不足以判断":
            add_cell_line(analysis_cell, "狼大分层", analysis["layer"])
        if analysis["trend_stage"] != "不足以判断":
            add_cell_line(analysis_cell, "趋势阶段", analysis["trend_stage"])
        if analysis["t_notes"]:
            add_cell_line(analysis_cell, "做T条件", "；".join(analysis["t_notes"]))
        if analysis["account_fit"] != "不足以判断":
            add_cell_line(analysis_cell, "账户适配", analysis["account_fit"])
        if analysis["refs"]:
            add_cell_line(analysis_cell, "图片框架补充", "；".join(analysis["refs"]))
        if analysis["action"] != "观察/未明确":
            add_cell_line(
                analysis_cell,
                "动作和风控",
                f"更接近“{analysis['action']}”。需结合原文触发条件、失效条件和当时盘面验证。",
            )
        if analysis["chain"]:
            add_cell_line(analysis_cell, "判断链", " / ".join(analysis["chain"]))
        if analysis["hypothesis"]:
            add_cell_line(analysis_cell, "可证伪假设", analysis["hypothesis"])
        if analysis["postmortem"]:
            add_cell_line(analysis_cell, "后验复盘", analysis["postmortem"])
        if analysis["memory"]:
            add_cell_line(analysis_cell, "交易记忆", analysis["memory"])
        if analysis["quality"]:
            add_cell_line(analysis_cell, "质检提示", analysis["quality"])
    doc.add_paragraph()


def add_day_toc(doc: Document, day_posts: list[dict[str, Any]]) -> None:
    doc.add_heading("当日发言目录", level=3)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["用户", "时间", "楼层", "类型", "方向提示"]
    for index, head in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, "E5E7EB")
        run = cell.paragraphs[0].add_run(head)
        set_font(run, 9, bold=True)
    for post in day_posts:
        cls, _ = classify_post(post)
        sectors = detect_sectors(post["body"])
        row = table.add_row().cells
        values = [
            post["username"],
            post["time"][11:16],
            f"{post['floor']}楼",
            f"{cls}类",
            "、".join(sectors[:3]) if sectors else "原文保留",
        ]
        for cell, value in zip(row, values):
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            run = cell.paragraphs[0].add_run(value)
            set_font(run, 8.5, color="374151")
    doc.add_paragraph()


def week_info(day: str) -> tuple[str, str]:
    current = datetime.strptime(day, "%Y-%m-%d")
    start = current - timedelta(days=current.weekday())
    end = start + timedelta(days=6)
    iso = current.isocalendar()
    key = f"{iso.year}-W{iso.week:02d}"
    label = f"第{iso.week:02d}周（{start:%Y-%m-%d} 至 {end:%Y-%m-%d}）"
    return key, label


def group_days_by_week(grouped: dict[str, list[dict[str, Any]]]) -> list[tuple[str, list[str]]]:
    week_days: dict[str, list[str]] = defaultdict(list)
    week_labels: dict[str, str] = {}
    for day in sorted(grouped):
        key, label = week_info(day)
        week_days[key].append(day)
        week_labels[key] = label
    return [(week_labels[key], week_days[key]) for key in sorted(week_days)]


def day_class_counts(day_posts: list[dict[str, Any]]) -> Counter[str]:
    counts: Counter[str] = Counter()
    for post in day_posts:
        cls, _ = classify_post(post)
        counts[cls] += 1
    return counts


def day_sector_hint(day_posts: list[dict[str, Any]]) -> str:
    counter: Counter[str] = Counter()
    for post in day_posts:
        cls, _ = classify_post(post)
        if cls != "C":
            counter.update(detect_sectors(post["body"]))
    hints = top_items(counter, 3)
    return "、".join(hints) if hints else "以原文保留为主"


def add_week_overview(doc: Document, days: list[str], grouped: dict[str, list[dict[str, Any]]]) -> None:
    doc.add_heading("本周阅读导航", level=2)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["日期", "发言数", "A/B/C", "主要方向", "阅读建议"]
    for index, head in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, "E5E7EB")
        run = cell.paragraphs[0].add_run(head)
        set_font(run, 9, bold=True)
    for day in days:
        day_posts = grouped[day]
        counts = day_class_counts(day_posts)
        valuable = counts["A"] + counts["B"]
        row = table.add_row().cells
        values = [
            day,
            str(len(day_posts)),
            f"A{counts['A']} / B{counts['B']} / C{counts['C']}",
            day_sector_hint(day_posts),
            "优先读A/B类与每日操盘内参" if valuable else "只看原文氛围和每日总结",
        ]
        for cell, value in zip(row, values):
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            run = cell.paragraphs[0].add_run(value)
            set_font(run, 8.5, color="374151")
    doc.add_paragraph()


def markdown_week_overview(days: list[str], grouped: dict[str, list[dict[str, Any]]]) -> list[str]:
    lines = ["### 本周阅读导航", "", "| 日期 | 发言数 | A/B/C | 主要方向 | 阅读建议 |", "| --- | ---: | --- | --- | --- |"]
    for day in days:
        day_posts = grouped[day]
        counts = day_class_counts(day_posts)
        valuable = counts["A"] + counts["B"]
        advice = "优先读A/B类与每日操盘内参" if valuable else "只看原文氛围和每日总结"
        lines.append(f"| {day} | {len(day_posts)} | A{counts['A']} / B{counts['B']} / C{counts['C']} | {day_sector_hint(day_posts)} | {advice} |")
    lines.append("")
    return lines


def report_label_for_posts(posts: list[dict[str, Any]], explicit_label: str | None = None) -> str:
    if explicit_label:
        return explicit_label
    month = posts[0]["posted_at"][:7]
    return f"{month[:4]}年{int(month[5:7])}月"


def build_doc(posts: list[dict[str, Any]], market_days: dict[str, MarketDay], resolver: StockResolver, out_path: Path, report_label: str | None = None) -> None:
    doc = Document()
    style_document(doc)
    month_label = report_label_for_posts(posts, report_label)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(f"{THREAD_TITLE}：{month_label}发言逐条分析")

    wolf_count = sum(1 for p in posts if p["uid"] == WOLF_UID)
    summary = doc.add_paragraph()
    summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = summary.add_run(f"共 {len(posts)} 条；狼大 {wolf_count} 条重点详析；其他 A/B 类发言按同框架分析；按发言时间顺序排列")
    set_font(r, 10, color="6B7280")

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = note.add_run(framework_note())
    set_font(nr, 8.8, color="6B7280")

    doc.add_heading("用户发言目录", level=1)
    counts = Counter((p["uid"], p["username"]) for p in posts)
    overview = doc.add_table(rows=1, cols=3)
    overview.style = "Table Grid"
    for i, head in enumerate(["用户", "UID", "本月发言数"]):
        cell = overview.rows[0].cells[i]
        set_cell_shading(cell, "E5E7EB")
        rr = cell.paragraphs[0].add_run(head)
        set_font(rr, 9.5, bold=True)
    for (uid, name), count in counts.most_common():
        row = overview.add_row().cells
        row[0].text = name
        row[1].text = uid
        row[2].text = str(count)
    doc.add_paragraph()

    grouped: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for post in posts:
        grouped[post["posted_at"][:10]].append(post)

    first_week = True
    for week_label, days in group_days_by_week(grouped):
        if not first_week:
            doc.add_section(WD_SECTION_START.NEW_PAGE)
        first_week = False
        week_total = sum(len(grouped[day]) for day in days)
        doc.add_heading(f"{week_label}（{week_total}条）", level=1)
        add_week_overview(doc, days, grouped)
        for day in days:
            day_posts = grouped[day]
            doc.add_heading(f"{day}（{len(day_posts)}条）", level=2)
            p = doc.add_paragraph()
            rr = p.add_run("当日指数环境：")
            set_font(rr, 10, bold=True)
            vv = p.add_run(market_summary(day, market_days))
            set_font(vv, 10)
            add_day_toc(doc, day_posts)

            for post in day_posts:
                analysis = make_analysis(post, market_days, resolver)
                add_post_card(doc, post, analysis)
            add_day_summary(doc, day, day_posts, market_days, resolver)

    doc.save(out_path)


def write_markdown(posts: list[dict[str, Any]], market_days: dict[str, MarketDay], resolver: StockResolver, out_path: Path, report_label: str | None = None) -> None:
    month_label = report_label_for_posts(posts, report_label)
    lines = [f"# {THREAD_TITLE}：{month_label}发言逐条分析", "", f"> {framework_note()}", ""]
    grouped: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for post in posts:
        grouped[post["posted_at"][:10]].append(post)
    for week_label, days in group_days_by_week(grouped):
        week_total = sum(len(grouped[day]) for day in days)
        lines.append(f"## {week_label}（{week_total}条）")
        lines.append("")
        lines.extend(markdown_week_overview(days, grouped))
        for day in days:
            day_posts = grouped[day]
            lines.append(f"### {day}（{len(day_posts)}条）")
            lines.append("")
            lines.append(f"当日指数环境：{market_summary(day, market_days)}")
            lines.append("")
            lines.append("#### 当日发言目录")
            lines.append("")
            lines.append("| 用户 | 时间 | 楼层 | 类型 | 方向提示 |")
            lines.append("| --- | --- | --- | --- | --- |")
            for post in day_posts:
                cls, _ = classify_post(post)
                sectors = detect_sectors(post["body"])
                hint = "、".join(sectors[:3]) if sectors else "原文保留"
                lines.append(f"| {post['username']} | {post['time'][11:16]} | {post['floor']}楼 | {cls}类 | {hint} |")
            lines.append("")
            for post in day_posts:
                analysis = make_analysis(post, market_days, resolver)
                lines.append(f"#### {post['username']} | {post['time']} | {post['floor']}楼 | UID {post['uid']} | {analysis['class']}类")
                lines.append("")
                lines.append(post["body"])
                lines.append("")
                if analysis["class"] != "C":
                    lines.append(f"**{'狼大发言详析' if post['uid'] == WOLF_UID else '发言分析'}**")
                    lines.append("")
                    lines.append(f"- 一句话还原：{analysis['translation']}")
                    if analysis["sectors"]:
                        lines.append(f"- 板块/细分方向：{'、'.join(analysis['sectors'])}")
                    if analysis["stocks"]:
                        lines.append("- 个股简称补全：" + "；".join(f"{s['alias']} → {s['name']}（{s['code']}，{s['direction']}，{s['source']}）" for s in analysis["stocks"]))
                    if analysis["referenced_markets"]:
                        lines.append("- 跨日市场环境：" + "；".join(analysis["referenced_markets"]))
                    if analysis["terms"]:
                        lines.append("- 术语解释：" + "；".join(f"{term}：{desc}" for term, desc in analysis["terms"]))
                    if analysis["tech"]:
                        lines.append("- 技术/量能结构：" + "；".join(analysis["tech"]))
                    if analysis["layer"] != "不足以判断":
                        lines.append(f"- 狼大分层：{analysis['layer']}")
                    if analysis["trend_stage"] != "不足以判断":
                        lines.append(f"- 趋势阶段：{analysis['trend_stage']}")
                    if analysis["t_notes"]:
                        lines.append("- 做T条件：" + "；".join(analysis["t_notes"]))
                    if analysis["account_fit"] != "不足以判断":
                        lines.append(f"- 账户适配：{analysis['account_fit']}")
                    if analysis["refs"]:
                        lines.append("- 图片框架补充：" + "；".join(analysis["refs"]))
                    if analysis["action"] != "观察/未明确":
                        lines.append(f"- 动作和风控：更接近“{analysis['action']}”。需结合原文触发条件、失效条件和当时盘面验证。")
                    if analysis["chain"]:
                        lines.append("- 判断链还原：" + " / ".join(analysis["chain"]))
                    if analysis["hypothesis"]:
                        lines.append(f"- 可证伪假设：{analysis['hypothesis']}")
                    if analysis["postmortem"]:
                        lines.append(f"- 后验复盘：{analysis['postmortem']}")
                    if analysis["memory"]:
                        lines.append(f"- 交易记忆：{analysis['memory']}")
                    if analysis["quality"]:
                        lines.append(f"- 质检提示：{analysis['quality']}")
                lines.append("")
            lines.append("#### 每日操盘内参")
            lines.append("")
            lines.append(f"> {day_summary_reference_note()}")
            lines.append("")
            for index, (head, value) in enumerate(build_day_summary(day, day_posts, market_days, resolver).items(), start=1):
                lines.append(f"##### {index}. {head}")
                lines.append("")
                lines.append(value)
                lines.append("")
    out_path.write_text("\n".join(lines), encoding="utf-8")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build analyzed June document.")
    parser.add_argument("--input", default="selected_users_posts.json")
    parser.add_argument("--month", default="2026-06")
    parser.add_argument("--start-date", help="Optional first date to include, e.g. 2026-06-22.")
    parser.add_argument("--end-date", help="Optional last date to include, e.g. 2026-06-28.")
    parser.add_argument("--out", default="monthly_docs/科学技术打头阵_发言逐条分析_2026-06.docx")
    parser.add_argument("--markdown", default="monthly_docs/科学技术打头阵_发言逐条分析_2026-06.md")
    parser.add_argument("--market-cache", default="market_cache_2026-06.json")
    parser.add_argument("--stock-cache", default="stock_name_cache.json")
    parser.add_argument("--label", help="Optional report label, e.g. 2026-W26（2026-06-22 至 2026-06-28）.")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    posts = load_posts(Path(args.input), args.month, args.start_date, args.end_date)
    if not posts:
        print("No posts found for selected period.")
        return 1
    if args.start_date and args.end_date:
        begin = args.start_date.replace("-", "")
        end = args.end_date.replace("-", "")
    else:
        begin = args.month.replace("-", "") + "01"
        year, month = map(int, args.month.split("-"))
        end = f"{year}{month:02d}{calendar.monthrange(year, month)[1]:02d}"
    market_days = fetch_index_klines(begin, end, Path(args.market_cache))
    resolver = StockResolver(Path(args.stock_cache))
    out = Path(args.out)
    out.parent.mkdir(parents=True, exist_ok=True)
    build_doc(posts, market_days, resolver, out, args.label)
    write_markdown(posts, market_days, resolver, Path(args.markdown), args.label)
    resolver.save()
    print(f"Generated {out} with {len(posts)} posts.")
    print(f"Generated {args.markdown}.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
